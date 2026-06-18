# -*- coding: utf-8 -*-
"""生成编译器项目第二周工作周报 Word 文档 — 语法分析"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)


def add_code_block(doc, code_text, font_size=9):
    """添加等宽字体代码块"""
    p = doc.add_paragraph()
    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(font_size)
    return p


def add_table_with_header(doc, headers, data, col_widths=None):
    """添加带表头的表格"""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    # 数据行
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table


# ══════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('编译原理课程设计')
run.font.size = Pt(28)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('第二周工作周报')
run2.font.size = Pt(22)
run2.font.name = '黑体'
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub2.add_run('—— 语法分析器的设计与实现')
run3.font.size = Pt(16)
run3.font.name = '宋体'
run3.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

for text in [
    '项目名称：Micro-C 编译器',
    '报告周期：2026年6月12日 — 2026年6月18日',
    '报告人：___（填写）',
    '指导教师：___（填写）',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ── 目录占位 ──
doc.add_heading('目  录', level=1)
doc.add_paragraph('（请在 Word 中插入自动目录：引用 → 目录 → 自动目录）')
doc.add_page_break()

# ══════════════════════════════════════════════
# 一、本周工作概述
# ══════════════════════════════════════════════
doc.add_heading('一、本周工作概述', level=1)

doc.add_paragraph(
    '本周是编译原理课程设计的第二周，承接上周完成的词法分析器（Lexer），'
    '本周的核心任务是完成编译器前端的第二个关键阶段——语法分析器（Parser）的设计与实现。'
    '围绕这一核心目标，本周的工作涵盖了以下几个主要方面：'
)

doc.add_heading('1.1 文法定义与理论基础', level=2)
doc.add_paragraph(
    '首先，根据 Micro-C 语言的语法特性，设计了完整的上下文无关文法（CFG）。'
    '本项目采用算符优先分析法（Operator Precedence Parsing）作为语法分析的核心算法，'
    '这是一种自底向上的移进-归约分析方法，特别适合处理表达式的优先级和结合性问题。'
    '在文法设计阶段，需要确保文法满足算符优先分析的约束条件：'
    '不含空产生式（ε）、不含相同右部、不含相邻非终结符。'
    '为此，对 Micro-C 的语法结构进行了严格的文法改造，最终定义了 28 条产生式，'
    '覆盖程序结构、声明语句、赋值语句、函数调用、控制流语句和表达式等全部语法结构。'
)

doc.add_heading('1.2 FIRSTVT/LASTVT 集合计算与优先关系表构建', level=2)
doc.add_paragraph(
    '在文法定义的基础上，实现了 FIRSTVT（最左终结符集）和 LASTVT（最右终结符集）的自动计算。'
    '这两个集合是构建优先关系表的关键前置条件。'
    'FIRSTVT 集合表示每个非终结符在推导过程中最左边可能出现的终结符，'
    'LASTVT 集合则表示最右边可能出现的终结符。'
    '通过不动点迭代算法，反复计算直到集合不再变化，确保了计算的完备性。'
    '随后，基于 FIRSTVT 和 LASTVT 集合，自动构建了终结符之间的优先关系表。'
    '表中定义了三种优先关系：移进（<）、等优先（=）和归约（>），'
    '为语法分析器的运行时决策提供了静态查表依据。'
)

doc.add_heading('1.3 算符优先语法分析器的实现', level=2)
doc.add_paragraph(
    '在优先关系表的基础上，实现了完整的算符优先语法分析器。'
    '分析器采用移进-归约的方式工作：维护一个符号栈，根据栈顶终结符与当前输入终结符的优先关系，'
    '决定是移进（将输入符号压栈）还是归约（将栈顶的产生式右部替换为左部非终结符）。'
    '针对 Micro-C 语言的特殊语法结构（如成员访问 id.id、数组下标 id[E]、赋值语句的分号处理等），'
    '在标准算法的基础上加入了大量的特殊处理逻辑和前瞻判断，'
    '使分析器能够正确处理 C-like 语言的复杂语法结构。'
)

doc.add_heading('1.4 语义分析与代码生成', level=2)
doc.add_paragraph(
    '在语法分析的同时，实现了语法制导翻译（Syntax-Directed Translation）。'
    '每当归约成功时，根据产生式的模式执行对应的语义动作，包括：'
    '符号表管理（变量声明、结构体定义）、类型检查（赋值兼容性、成员访问合法性）、'
    '中间代码生成（四元式）和控制流的标签回填。'
    '最终，四元式序列被翻译为 Jack 虚拟机指令集的目标代码，完成了从源代码到目标代码的完整编译流程。'
)

doc.add_heading('1.5 测试与验证', level=2)
doc.add_paragraph(
    '为了验证语法分析器的正确性，编写了覆盖所有语言特性的综合测试用例。'
    '测试用例包含变量声明、结构体定义与成员访问、赋值语句、算术与关系表达式、'
    '输入输出语句、循环语句、选择语句和返回语句等全部 Micro-C 支持的语法结构。'
    '通过运行编译器并检查输出的移进-归约跟踪、语法分析树、四元式序列和目标代码，'
    '确认了各模块的功能正确性。同时，也测试了一些包含语法错误的源文件，验证了错误处理机制的有效性。'
)

# ══════════════════════════════════════════════
# 二、详细技术实现
# ══════════════════════════════════════════════
doc.add_heading('二、详细技术实现', level=1)

doc.add_heading('2.1 项目目录结构', level=2)
doc.add_paragraph('本周在上周项目结构的基础上，新增了语法分析和语义分析相关的模块，完整的项目目录结构如下：')

add_code_block(doc,
    'ShiXvn_Compiler/\n'
    '├── main.cpp                        # 程序入口（三阶段串联）\n'
    '├── test.txt                        # 综合测试源文件\n'
    '├── include/\n'
    '│   ├── Token.h                     # Token 数据结构定义\n'
    '│   ├── ErrorHandler.h              # 统一错误处理接口\n'
    '│   ├── SymbolTable.h               # 词法层符号表\n'
    '│   ├── Lexer.h                     # 词法分析器接口\n'
    '│   ├── Grammar.h                   # ★ 文法定义与优先关系表\n'
    '│   ├── PrecedenceParser.h          # ★ 算符优先语法分析器\n'
    '│   └── SemanticAnalyzer.h          # ★ 语义分析器与代码生成\n'
    '├── src/\n'
    '│   ├── ErrorHandler.cpp\n'
    '│   ├── SymbolTable.cpp\n'
    '│   ├── Lexer.cpp\n'
    '│   ├── Grammar.cpp                 # ★ 文法实现 + FIRSTVT/LASTVT/优先表\n'
    '│   ├── PrecedenceParser.cpp        # ★ 移进-归约主控 + 语义动作\n'
    '│   └── SemanticAnalyzer.cpp        # ★ 语义检查 + 四元式 + Jack VM\n'
    '└── output/\n'
    '    ├── Lexer/tokens.txt            # 词法分析输出\n'
    '    ├── Parser/                     # 语法分析输出\n'
    '    │   ├── productions.txt         #   产生式列表\n'
    '    │   ├── firstvt.txt             #   FIRSTVT 集合\n'
    '    │   ├── lastvt.txt              #   LASTVT 集合\n'
    '    │   ├── precedence_table.txt    #   优先关系表\n'
    '    │   ├── trace.txt               #   移进-归约跟踪\n'
    '    │   └── parse_tree.txt          #   语法分析树\n'
    '    └── Codegen/                    # 代码生成输出\n'
    '        ├── quads.txt               #   四元式序列\n'
    '        └── target.vm               #   Jack VM 目标代码'
)

doc.add_paragraph()
doc.add_paragraph(
    '注：标有 ★ 的文件为本周新增模块，是语法分析阶段的核心组成部分。'
    '与上周相比，新增了 Grammar、PrecedenceParser、SemanticAnalyzer 三个核心类，'
    '分别负责文法定义、语法分析驱动和语义分析与代码生成。'
)

doc.add_heading('2.2 文法定义（Grammar 模块）', level=2)

doc.add_heading('2.2.1 产生式定义', level=3)
doc.add_paragraph(
    'Micro-C 的文法共定义了 28 条产生式，覆盖了语言的所有语法结构。'
    '产生式采用编号管理，每条产生式包含左部非终结符和右部符号序列。'
    '下表列出了全部产生式及其分类：'
)

prod_data = [
    ['1', 'P → S ; P', '程序结构（递归）'],
    ['2', 'P → S ;', '程序结构（终止）'],
    ['3', 'S → int id', '整型变量声明'],
    ['4', 'S → char id', '字符型变量声明'],
    ['5', 'S → double id', '浮点型变量声明'],
    ['6', 'S → struct id { P }', '结构体类型定义'],
    ['7', 'S → struct id id', '结构体变量声明'],
    ['8', 'S → id = E', '简单赋值语句'],
    ['9', 'S → id . id = E', '结构体成员赋值'],
    ['10', 'S → id [ E ] = E', '数组元素赋值'],
    ['11', 'S → id ( E )', '函数调用（带参数）'],
    ['12', 'S → id ( id )', '函数调用（标识符参数）'],
    ['13', 'S → id ( )', '函数调用（无参数）'],
    ['14', 'S → while ( E ) { P }', '循环语句'],
    ['15', 'S → if ( E ) { P } else { P }', '选择语句'],
    ['16', 'S → return E', '返回语句'],
    ['17', 'E → E + E', '加法表达式'],
    ['18', 'E → E - E', '减法表达式'],
    ['19', 'E → E * E', '乘法表达式'],
    ['20', 'E → E / E', '除法表达式'],
    ['21', 'E → E < E', '小于关系表达式'],
    ['22', 'E → E > E', '大于关系表达式'],
    ['23', 'E → E == E', '等于关系表达式'],
    ['24', 'E → E . id', '结构体成员访问'],
    ['25', 'E → E [ E ]', '数组下标访问'],
    ['26', 'E → ( E )', '括号表达式'],
    ['27', 'E → id', '标识符（变量）'],
    ['28', 'E → num', '整型常数'],
    ['29', 'E → char', '字符常数'],
]
add_table_with_header(doc, ['编号', '产生式', '语义说明'], prod_data)
doc.add_paragraph()

doc.add_paragraph(
    '文法的设计遵循了算符优先分析的约束条件。其中，非终结符共 3 个：'
    'P（程序）、S（语句）、E（表达式）；终结符共 25 个，包括关键字、运算符、界符和基本元素。'
    '开始符号为 P，程序由若干分号分隔的语句构成。'
)

doc.add_heading('2.2.2 文法设计的关键考量', level=3)
doc.add_paragraph(
    '在文法设计过程中，需要特别注意以下几个关键问题：'
)
items = [
    '消除空产生式：Micro-C 的文法中不允许出现 ε 产生式，所有推导必须有实质性的符号。这要求在设计声明语句时，即使没有初始化表达式，也要保留类型关键字和标识符。',
    '消除相同右部：不同的非终结符不能推导出完全相同的结构，否则归约时会产生二义性。本项目中 P（程序）和 S（语句）的右部严格区分，P 负责语句序列的组织，S 负责单个语句的结构。',
    '避免相邻非终结符：产生式右部尽量用终结符将非终结符隔开，以便通过终结符来判断优先级。例如 P → S ; P 中，S 和 P 之间用分号隔开；S → id = E 中，id 和 E 之间用等号隔开。',
    '表达式的递归定义：算术和关系表达式采用 E → E op E 的形式定义，这天然地引入了左递归。算符优先分析法通过优先关系表来处理运算符的优先级和结合性，无需消除左递归。',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.3 FIRSTVT 与 LASTVT 集合计算', level=2)

doc.add_heading('2.3.1 算法原理', level=3)
doc.add_paragraph(
    'FIRSTVT(A) 表示非终结符 A 在推导过程中最左边可能出现的终结符集合。'
    '其计算规则如下：'
)
rules = [
    '若存在产生式 A → a... 或 A → B a ...（其中 a 为终结符，B 为非终结符），则 a ∈ FIRSTVT(A)；',
    '若存在产生式 A → B... 且 a ∈ FIRSTVT(B)，则 a ∈ FIRSTVT(A)（传播规则）。',
]
for r in rules:
    doc.add_paragraph(r, style='List Number')

doc.add_paragraph(
    'LASTVT(A) 的计算规则与 FIRSTVT 对称，从产生式右部的末尾开始分析：'
)
rules2 = [
    '若存在产生式 A → ...a 或 A → ...a B（其中 a 为终结符，B 为非终结符），则 a ∈ LASTVT(A)；',
    '若存在产生式 A → ...B 且 a ∈ LASTVT(B)，则 a ∈ LASTVT(A)（传播规则）。',
]
for r in rules2:
    doc.add_paragraph(r, style='List Number')

doc.add_heading('2.3.2 不动点迭代算法', level=3)
doc.add_paragraph(
    'FIRSTVT 和 LASTVT 集合的计算采用不动点迭代算法。'
    '算法从空集合开始，反复扫描所有产生式，根据上述规则向集合中添加元素。'
    '当某一轮扫描中没有任何集合发生变化时，算法终止，此时的集合即为最终结果。'
    '该算法保证了计算的完备性，不会遗漏任何可能的终结符。'
)

doc.add_paragraph('算法伪代码如下：')
add_code_block(doc,
    'function computeFirstVT(productions, nonTerminals):\n'
    '    for each A in nonTerminals:\n'
    '        firstVT[A] = {}\n'
    '    changed = true\n'
    '    while changed:\n'
    '        changed = false\n'
    '        for each production A → α in productions:\n'
    '            if α[0] is terminal:\n'
    '                if firstVT[A].add(α[0]): changed = true\n'
    '            if len(α) >= 2 and α[0] is nonTerminal and α[1] is terminal:\n'
    '                if firstVT[A].add(α[1]): changed = true\n'
    '            if α[0] is nonTerminal:\n'
    '                for each a in firstVT[α[0]]:\n'
    '                    if firstVT[A].add(a): changed = true'
)

doc.add_paragraph()

doc.add_heading('2.3.3 计算结果示例', level=3)
doc.add_paragraph('以下是部分非终结符的 FIRSTVT 和 LASTVT 集合计算结果：')

fv_data = [
    ['P', '{, int, char, double, struct, id, while, if, return', '{ ; }'],
    ['S', '{ int, char, double, struct, id, while, if, return', '{ id, ;, } }'],
    ['E', '{ (, id, num, char', '{ ), id, num, char, +, -, *, /, <, >, ==, ., [ }'],
]
add_table_with_header(doc, ['非终结符', 'FIRSTVT 集合', 'LASTVT 集合'], fv_data)
doc.add_paragraph()
doc.add_paragraph(
    '（完整的 FIRSTVT 和 LASTVT 集合已输出至 output/Parser/firstvt.txt 和 output/Parser/lastvt.txt 文件）'
)

doc.add_heading('2.4 优先关系表构建', level=2)

doc.add_heading('2.4.1 构建规则', level=3)
doc.add_paragraph('优先关系表的构建基于以下四条规则：')

table_rules = [
    ['规则1', '相邻终结符相等', '若产生式右部出现 Xi Xi+1（两个相邻终结符），则 Xi = Xi+1'],
    ['规则2', '终结符 < 非终结符', '若产生式右部出现 Xi B（终结符后跟非终结符），则 Xi < a，其中 a ∈ FIRSTVT(B)'],
    ['规则3', '非终结符 > 终结符', '若产生式右部出现 B Xi（非终结符后跟终结符），则 a > Xi，其中 a ∈ LASTVT(B)'],
    ['规则4', '栈底哨兵', '$ < a（a ∈ FIRSTVT(P)），a > $（a ∈ LASTVT(P)）'],
]
add_table_with_header(doc, ['规则', '名称', '描述'], table_rules)
doc.add_paragraph()

doc.add_paragraph(
    '此外，还需要手动添加一些特殊关系以处理括号匹配和赋值语句结束：'
    '( ) = 、[ ] = 、{ } = （配对括号等优先），以及 = > ; （赋值表达式后跟分号时归约）。'
)

doc.add_heading('2.4.2 优先关系表（局部展示）', level=3)
doc.add_paragraph('以下是优先关系表的部分内容（行表示栈顶终结符，列表示输入终结符）：')

# 一个简化的优先关系表
rel_headers = ['', 'id', 'num', '+', '*', '(', ')', '=', ';', '{', '}', '$']
rel_data = [
    ['id',  '',  '',  '>', '>', '',  '>', '>', '>', '',  '>', '>'],
    ['num', '',  '',  '>', '>', '',  '>', '>', '>', '',  '>', '>'],
    ['+',   '<', '<', '',  '',  '<', '>', '',  '>', '',  '',  '>'],
    ['*',   '<', '<', '',  '',  '<', '>', '',  '>', '',  '',  '>'],
    ['(',   '<', '<', '',  '',  '<', '=', '',  '',  '',  '',  ''],
    [')',   '',  '',  '>', '>', '',  '>', '>', '>', '',  '>', '>'],
    ['=',   '<', '<', '',  '',  '<', '',  '',  '>', '',  '',  '>'],
    [';',   '',  '',  '',  '',  '',  '',  '',  '',  '<',  '',  '>'],
    ['{',   '<', '<', '',  '',  '<', '',  '',  '<', '',  '=', ''],
    ['}',   '',  '',  '>', '>', '',  '>', '>', '>', '',  '>', '>'],
    ['$',   '<', '<', '',  '',  '<', '',  '',  '',  '<',  '',  '='],
]
add_table_with_header(doc, rel_headers, rel_data)
doc.add_paragraph()
doc.add_paragraph(
    '（完整的优先关系表已输出至 output/Parser/precedence_table.txt 文件，'
    '包含全部 25 个终结符之间的优先关系。）'
)

doc.add_heading('2.5 算符优先语法分析器（PrecedenceParser 模块）', level=2)

doc.add_heading('2.5.1 核心数据结构', level=3)
doc.add_paragraph('语法分析器的核心数据结构定义如下：')

add_code_block(doc,
    '// 分析栈中的元素\n'
    'struct StackSymbol {\n'
    '    string symbol;       // 符号名\n'
    '    bool isTerminal;     // 是否为终结符\n'
    '    int line;            // 源码行号\n'
    '    int treeNodeId;      // 对应语法树节点 ID\n'
    '    string addr;         // 语义值（变量名/临时变量/常量）\n'
    '    string type;         // 类型信息\n'
    '};\n'
    '\n'
    '// 语法树节点\n'
    'struct TreeNode {\n'
    '    int id;              // 节点 ID\n'
    '    string symbol;       // 符号名\n'
    '    bool isTerminal;     // 是否为终结符\n'
    '    vector<int> children;// 子节点 ID 列表\n'
    '    int line;            // 源码行号\n'
    '};\n'
    '\n'
    '// 移进-归约跟踪记录\n'
    'struct TraceEntry {\n'
    '    string stackStr;     // 当前栈内容\n'
    '    string inputStr;     // 剩余输入\n'
    '    string action;       // 执行的动作\n'
    '};'
)

doc.add_paragraph()
doc.add_paragraph(
    'StackSymbol 结构体是分析栈的基本元素，除了包含符号名和类型标记外，'
    '还携带了语义值（addr）和语法树节点 ID（treeNodeId），'
    '使得在移进和归约过程中能够同步进行语义分析和语法树构建。'
)

doc.add_heading('2.5.2 Token 到终结符的映射', level=3)
doc.add_paragraph(
    '词法分析器输出的 Token 需要映射为文法中的终结符符号，才能进行优先关系查表。'
    '映射规则如下：'
)

map_data = [
    ['KEYWORD', 'main, read, write', 'id', '作为函数名处理'],
    ['KEYWORD', '其他关键字', '原值', 'int, char, while, if 等保留'],
    ['IDENTIFIER', '-', 'id', '统一映射为 id'],
    ['NUMBER', '-', 'num', '统一映射为 num'],
    ['CHAR_LITERAL', '-', 'char', '统一映射为 char'],
    ['OPERATOR', '-', '原值', '+, -, *, /, =, . 等保留'],
    ['DELIMITER', '-', '原值', '(, ), {, }, ;, [, ] 保留'],
    ['END_OF_FILE', '-', '$', '文件结束标记'],
]
add_table_with_header(doc, ['Token 类型', '条件', '映射为', '说明'], map_data)
doc.add_paragraph()

doc.add_heading('2.5.3 移进-归约主控算法', level=3)
doc.add_paragraph(
    '语法分析器的核心是移进-归约主控循环。算法维护一个符号栈和一个输入流，'
    '根据栈顶终结符与当前输入终结符的优先关系，决定下一步动作。'
    '整体流程如下图所示：'
)

doc.add_paragraph()
doc.add_paragraph('[流程图描述：算符优先语法分析器主控循环流程图]')
doc.add_paragraph()

add_code_block(doc,
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│                      算符优先分析主控循环                        │\n'
    '├──────────────────────────────────────────────────────────────────┤\n'
    '│                                                                  │\n'
    '│   ┌─────────────┐                                               │\n'
    '│   │ 初始化：     │                                               │\n'
    '│   │ 栈 = [$]     │                                               │\n'
    '│   │ 输入 = Token流│                                              │\n'
    '│   └──────┬──────┘                                               │\n'
    '│          ▼                                                       │\n'
    '│   ┌──────────────┐     是     ┌─────────────┐                   │\n'
    '│   │ 栈=[$P] 且   │───────────→│   ACCEPT    │                   │\n'
    '│   │ 输入=[$] ?   │            └─────────────┘                   │\n'
    '│   └──────┬───────┘                                               │\n'
    '│          │ 否                                                    │\n'
    '│          ▼                                                       │\n'
    '│   ┌──────────────────┐                                          │\n'
    '│   │ 查找栈顶终结符 a  │                                          │\n'
    '│   │ 查询关系 R(a, b)  │  b = 当前输入终结符                      │\n'
    '│   └──────┬───────────┘                                          │\n'
    '│          │                                                       │\n'
    '│    ┌─────┼─────┬─────────┐                                      │\n'
    '│    │     │     │         │                                      │\n'
    '│    ▼     ▼     ▼         ▼                                      │\n'
    '│  R="<" R="=" R=">"    R=""                                      │\n'
    '│  移进  移进  归约    报错/特殊处理                               │\n'
    '│    │     │     │         │                                      │\n'
    '│    └─────┴─────┴─────────┘                                      │\n'
    '│               │                                                  │\n'
    '│               ▼                                                  │\n'
    '│         继续主循环                                                │\n'
    '└──────────────────────────────────────────────────────────────────┘'
)

doc.add_paragraph()

doc.add_heading('2.5.4 标准移进操作', level=3)
doc.add_paragraph(
    '当优先关系为 <（移进）或 =（等优先）时，执行移进操作。'
    '移进操作将当前输入 Token 压入分析栈，同时为该符号创建语法树节点。'
    '对于终结符（如 id、num、char），其语义值（addr）被设置为 Token 的字面值，'
    '供后续语义动作使用。'
)

add_code_block(doc,
    '// 移进操作伪代码\n'
    'if relation == "<" or relation == "=":\n'
    '    node = createNode(currentTerminal, isTerminal=true, line)\n'
    '    sym = StackSymbol {\n'
    '        symbol: currentTerminal,\n'
    '        isTerminal: true,\n'
    '        line: currentToken.line,\n'
    '        treeNodeId: node.id,\n'
    '        addr: currentToken.value   // 语义值\n'
    '    }\n'
    '    stack.push(sym)\n'
    '    inputPos++'
)

doc.add_paragraph()

doc.add_heading('2.5.5 标准归约操作', level=3)
doc.add_paragraph(
    '当优先关系为 >（归约）时，执行归约操作。'
    '归约操作从栈顶向下搜索与某条产生式右部匹配的子串，'
    '找到后执行以下步骤：'
)

reduce_steps = [
    '执行语义动作（semanticAction）：根据产生式类型进行符号表操作、类型检查和四元式生成；',
    '收集子节点 ID：将被归约的符号对应的语法树节点收集为子节点列表；',
    '创建归约节点：为产生式左部非终结符创建新的语法树节点，并将子节点列表挂载到该节点下；',
    '栈操作：弹出产生式右部的所有符号，压入左部非终结符，保留语义值；',
    '记录跟踪日志：将归约操作记录到跟踪序列中，用于调试和验证。',
]
for s in reduce_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('归约操作的伪代码如下：')
add_code_block(doc,
    'function reduce():\n'
    '    for len = stack.size() downto 1:\n'
    '        start = stack.size() - len\n'
    '        candidate = stack[start .. top]\n'
    '        for each production A → α in grammar:\n'
    '            if α matches candidate:\n'
    '                // 1. 执行语义动作\n'
    '                semanticAction(production, start)\n'
    '                // 2. 收集子节点\n'
    '                childIds = [sym.treeNodeId for sym in stack[start..top]]\n'
    '                // 3. 创建归约节点\n'
    '                node = createNode(A, isTerminal=false)\n'
    '                node.children = childIds\n'
    '                // 4. 栈操作\n'
    '                addr = stack[start].addr\n'
    '                pop len symbols from stack\n'
    '                push StackSymbol(A, addr=addr)\n'
    '                // 5. 记录日志\n'
    '                trace.push(stack, "REDUCE: α → A")\n'
    '                return true\n'
    '    return false'
)

doc.add_paragraph()

doc.add_heading('2.5.6 特殊处理机制', level=3)
doc.add_paragraph(
    '标准的算符优先分析法无法直接处理 Micro-C 中的一些复杂语法结构。'
    '为此，在主控循环中加入了以下特殊处理机制，这些是本项目的核心创新点：'
)

spec_data = [
    ['成员访问/赋值\nid.id 模式', '当栈顶为 id . id 时，根据下一个 Token 的类型决定动作：\n'
     '- 下一个是 = → 移进（成员赋值 id.id=E）\n'
     '- 其他 → 归约 id.id → E（成员访问）'],
    ['赋值语句的\n分号处理', '当关系为空且输入为 ; 时，在栈中向下搜索 id=E 模式进行强制归约。\n'
     '解决了 = 和 ; 之间可能无优先关系的问题。'],
    ['数组下标处理\nid[E] 模式', '- 栈顶为 id 且输入为 [ 时，不归约 id→E，而是移进 [\n'
     '- 栈顶为 [ E 且输入为 ] 时，移进 ]\n'
     '- 栈顶为 [ E ] 时，归约为 E（数组下标表达式）'],
    ['数组元素赋值\nid[E]=E 模式', '当栈顶为 [ E ] 且 [ 前为 id 且输入为 = 时，移进 =，\n'
     '以支持数组元素赋值语法。'],
    ['结构体变量声明\nstruct id id', '在 > 关系下，若栈顶为 struct id 且输入为 id，\n'
     '则特殊移进以支持 struct Point p 声明语法。'],
    ['成员访问的\n前瞻判断', '在 > 关系下，若栈顶为 id 且当前为 . 且后面跟 id，\n'
     '则移进 .（而非归约 id→E），保留 id.id 的完整结构。'],
    ['结构体上下文\n检测', '当检测到 struct id { 模式时，设置结构体上下文标志，\n'
     '使后续的声明语句（如 int x）被识别为结构体成员定义。'],
]
add_table_with_header(doc, ['特殊处理', '详细说明'], spec_data)
doc.add_paragraph()

doc.add_heading('2.5.7 移进-归约过程示例', level=3)
doc.add_paragraph(
    '以下是对简单赋值语句 "x = 10 ;" 的移进-归约过程跟踪（部分步骤）：'
)

trace_data = [
    ['1', '$', 'id = 10 ; $', 'SHIFT: id'],
    ['2', '$ id', '= 10 ; $', 'SHIFT: ='],
    ['3', '$ id =', '10 ; $', 'SHIFT: num'],
    ['4', '$ id = num', '; $', 'REDUCE: num → E (prod 28)'],
    ['5', '$ id = E', '; $', 'REDUCE: id = E → S (prod 8)'],
    ['6', '$ S', '; $', 'SHIFT: ;'],
    ['7', '$ S ;', '$', 'REDUCE: S ; → P (prod 2)'],
    ['8', '$ P', '$', 'ACCEPT'],
]
add_table_with_header(doc, ['步骤', '栈', '输入', '动作'], trace_data)
doc.add_paragraph()
doc.add_paragraph(
    '（完整的移进-归约过程已输出至 output/Parser/trace.txt 文件）'
)

doc.add_heading('2.6 语法制导翻译（SemanticAnalyzer 模块）', level=2)

doc.add_heading('2.6.1 语义动作调度', level=3)
doc.add_paragraph(
    '每当归约成功时，根据产生式的右部模式（长度和符号序列）分派到对应的语义处理逻辑。'
    '语义动作调度器覆盖了以下所有语义场景：'
)

sem_data = [
    ['声明语句', 'int id / char id / double id / struct id id', '调用 declareVariable() 注册变量；\n在结构体体内时调用 addStructMember()'],
    ['结构体定义', 'struct id { P }', '调用 defineStruct() 注册类型；\n清除结构体上下文标志'],
    ['简单赋值', 'id = E', '类型检查 checkAssignment()；\n生成四元式 (=, exprAddr, "", varName)'],
    ['成员赋值', 'id . id = E', '成员访问检查 checkMemberAccess()；\n生成四元式 (=, exprAddr, "", "base.member")'],
    ['数组赋值', 'id [ E ] = E', '生成四元式 (=, exprAddr, "", "arrName[index]")'],
    ['函数调用', 'id(E) / id(id) / id()', '区分 read/write 内置函数；\n生成 param + call 四元式'],
    ['循环语句', 'while(E){P}', '生成标签 L0/L1；\n重排四元式：label L0 → 条件 → if_false → goto L0 → label L1'],
    ['选择语句', 'if(E){P}else{P}', '生成 if_false + goto；\n回填标签到跳转指令'],
    ['返回语句', 'return E', '生成四元式 (= return, retAddr, "", "")'],
    ['算术表达式', 'E op E', '生成临时变量 tN；\n生成四元式 (op, left, right, temp)'],
    ['成员访问', 'E . id', '生成临时变量；\n生成四元式 (., base, member, temp)'],
    ['数组下标', 'E [ E ]', '生成临时变量；\n生成四元式 ([], base, index, temp)'],
    ['括号表达式', '( E )', '传递子表达式的语义值'],
    ['基本元素', 'id / num / char', 'id 检查是否已声明（未声明则隐式声明为 int）；\nnum/char 的语义值在移进时已设置'],
]
add_table_with_header(doc, ['语义场景', '产生式模式', '语义动作'], sem_data)
doc.add_paragraph()

doc.add_heading('2.6.2 四元式生成', level=3)
doc.add_paragraph(
    '四元式（Quadruple）是本项目采用的中间代码表示形式。'
    '每条四元式包含四个字段：操作符（op）、第一操作数（arg1）、第二操作数（arg2）和结果（result）。'
    '下表列出了四元式中各操作符的含义：'
)

quad_data = [
    ['=', '赋值', '源操作数', '-', '目标变量'],
    ['+', '加法', '左操作数', '右操作数', '临时变量'],
    ['-', '减法', '左操作数', '右操作数', '临时变量'],
    ['*', '乘法', '左操作数', '右操作数', '临时变量'],
    ['/', '除法', '左操作数', '右操作数', '临时变量'],
    ['<', '小于', '左操作数', '右操作数', '临时变量'],
    ['>', '大于', '左操作数', '右操作数', '临时变量'],
    ['==', '等于', '左操作数', '右操作数', '临时变量'],
    ['.', '成员访问', '基对象', '成员名', '临时变量'],
    ['[]', '数组下标', '数组名', '下标', '临时变量'],
    ['label', '定义标签', '标签名', '-', '-'],
    ['goto', '无条件跳转', '-', '-', '目标标签'],
    ['if_false', '条件跳转', '条件变量', '-', '目标标签'],
    ['param', '参数传递', '参数值', '-', '-'],
    ['call', '函数调用', '函数名', '参数个数', '返回值'],
    ['return', '返回', '返回值', '-', '-'],
]
add_table_with_header(doc, ['操作符', '含义', 'ARG1', 'ARG2', 'RESULT'], quad_data)
doc.add_paragraph()

doc.add_heading('2.6.3 四元式示例', level=3)
doc.add_paragraph(
    '以下是对测试用例中部分语句生成的四元式序列（部分）：'
)

example_quad = [
    ['0', '=', '10', '-', 'x'],
    ['1', '=', '20', '-', 'y'],
    ['2', '=', '100', '-', 'p.px'],
    ['3', '=', '200', '-', 'p.py'],
    ['4', '+', 'x', 'y', 't0'],
    ['5', '=', 't0', '-', 'z'],
    ['6', '*', 'y', '2', 't1'],
    ['7', '+', 'x', 't1', 't2'],
    ['8', '=', 't2', '-', 'z'],
    ['9', 'param', 'z', '-', '-'],
    ['10', 'call', 'write', '1', '_'],
    ['11', 'call', 'println', '0', '_'],
    ['12', '<', 'x', '10', 't3'],
    ['13', 'if_false', 't3', '-', 'L1'],
    ['14', '+', 'x', '1', 't4'],
    ['15', '=', 't4', '-', 'x'],
    ['16', 'goto', '-', '-', 'L0'],
    ['17', 'label', 'L1', '-', '-'],
]
add_table_with_header(doc, ['ID', 'OP', 'ARG1', 'ARG2', 'RESULT'], example_quad)
doc.add_paragraph()
doc.add_paragraph('（完整的四元式序列已输出至 output/Codegen/quads.txt 文件）')

doc.add_heading('2.6.4 控制流的标签回填', level=3)
doc.add_paragraph(
    '控制流语句（while 循环和 if-else 选择）的四元式生成需要解决一个关键问题：'
    '在生成条件跳转指令时，跳转目标标签尚未确定（因为循环体或分支体还未生成）。'
    '解决方案是采用标签回填（Backpatching）技术：'
)

bp_steps = [
    '先生成跳转指令，但目标标签字段留空，记录该四元式的 ID；',
    '继续生成循环体或分支体的四元式；',
    '当循环体或分支体生成完毕后，确定目标标签的值；',
    '调用 backpatch() 函数，将标签值回填到之前留空的跳转指令中。',
]
for s in bp_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('以 while 循环为例，四元式的生成和回填过程如下：')

add_code_block(doc,
    'while (x < 10) { x = x + 1 ; }\n'
    '\n'
    '生成过程：\n'
    '  (1) 提取条件表达式四元式: (<, x, 10, t3)\n'
    '  (2) 生成 label L0 (循环开始)\n'
    '  (3) 恢复条件四元式\n'
    '  (4) 生成 if_false t3, ? (目标标签待回填) → 记录 ID = 13\n'
    '  (5) 生成循环体: (+, x, 1, t4), (=, t4, -, x)\n'
    '  (6) 生成 goto L0 (跳回循环开始)\n'
    '  (7) 生成 label L1 (循环结束)\n'
    '  (8) backpatch(13, L1) → 回填跳转目标\n'
    '\n'
    '最终四元式序列：\n'
    '  label L0\n'
    '  (<, x, 10, t3)\n'
    '  if_false t3, L1      ← 跳转目标已回填\n'
    '  (+, x, 1, t4)\n'
    '  (=, t4, -, x)\n'
    '  goto L0\n'
    '  label L1'
)

doc.add_paragraph()

doc.add_heading('2.7 目标代码生成', level=2)

doc.add_heading('2.7.1 Jack 虚拟机指令集', level=3)
doc.add_paragraph(
    '本项目将四元式序列翻译为 Jack 虚拟机（Jack VM）指令集的目标代码。'
    'Jack VM 是一种基于栈的虚拟机，其指令集包括：'
)

vm_data = [
    ['push segment index', '将指定位置的值压入栈顶'],
    ['pop segment index', '将栈顶值弹出到指定位置'],
    ['add / sub', '将栈顶两个值相加/相减，结果压入栈'],
    ['lt / gt / eq', '比较栈顶两个值，结果压入栈（布尔值）'],
    ['label labelName', '定义标签'],
    ['goto labelName', '无条件跳转到标签'],
    ['if-goto labelName', '条件跳转：栈顶为真时跳转'],
    ['call funcName nArgs', '调用函数，传递 nArgs 个参数'],
    ['return', '从函数返回，返回值在栈顶'],
    ['function name nLocals', '声明函数，分配 nLocals 个局部变量'],
    ['not', '对栈顶值取反'],
]
add_table_with_header(doc, ['指令', '说明'], vm_data)
doc.add_paragraph()

doc.add_heading('2.7.2 翻译规则', level=3)
doc.add_paragraph('四元式到 Jack VM 指令的翻译规则如下：')

trans_data = [
    ['=', 'push arg1; pop result'],
    ['+', 'push arg1; push arg2; add; pop result'],
    ['-', 'push arg1; push arg2; sub; pop result'],
    ['*', 'push arg1; push arg2; call Math.multiply 2; pop result'],
    ['/', 'push arg1; push arg2; call Math.divide 2; pop result'],
    ['<', 'push arg1; push arg2; lt; pop result'],
    ['>', 'push arg1; push arg2; gt; pop result'],
    ['==', 'push arg1; push arg2; eq; pop result'],
    ['label', 'label arg1'],
    ['goto', 'goto result'],
    ['if_false', 'push arg1; not; if-goto result'],
    ['param', 'push arg1'],
    ['call', 'call funcName nArgs; pop result（或 pop temp 0）'],
    ['return', 'push arg1; return'],
]
add_table_with_header(doc, ['四元式 OP', 'Jack VM 指令序列'], trans_data)
doc.add_paragraph()

doc.add_heading('2.7.3 内置函数映射', level=3)
doc.add_paragraph('Micro-C 的内置函数被映射为 Jack OS 的标准库函数：')

builtin_data = [
    ['read()', 'Keyboard.readInt', '从键盘读取一个整数'],
    ['write()', 'Output.printInt', '输出一个整数到屏幕'],
    ['println()', 'Output.println', '输出换行符'],
]
add_table_with_header(doc, ['Micro-C 函数', 'Jack OS 函数', '功能说明'], builtin_data)
doc.add_paragraph()

doc.add_heading('2.7.4 目标代码示例', level=3)
doc.add_paragraph('以下是测试用例中部分语句对应的目标代码（Jack VM 指令）：')

add_code_block(doc,
    '// Micro-C 源代码:\n'
    '// int x ;\n'
    '// int y ;\n'
    '// x = 10 ;\n'
    '// y = 20 ;\n'
    '// z = x + y * 2 ;\n'
    '// write(z) ;\n'
    '\n'
    'function Main.main 5          // 声明函数，5 个局部变量\n'
    '\n'
    '// x = 10\n'
    'push constant 10\n'
    'pop local 0\n'
    '\n'
    '// y = 20\n'
    'push constant 20\n'
    'pop local 1\n'
    '\n'
    '// z = x + y * 2\n'
    'push local 1                  // y\n'
    'push constant 2               // 2\n'
    'call Math.multiply 2          // y * 2\n'
    'pop temp 0\n'
    'push local 0                  // x\n'
    'push temp 0                   // y * 2\n'
    'add                           // x + y * 2\n'
    'pop local 2                   // z\n'
    '\n'
    '// write(z)\n'
    'push local 2\n'
    'call Output.printInt 1\n'
    'pop temp 0\n'
    'call Output.println 0\n'
    'pop temp 0\n'
    '\n'
    '// return 0\n'
    'push constant 0\n'
    'return'
)

doc.add_paragraph()
doc.add_paragraph('（完整的目标代码已输出至 output/Codegen/target.vm 文件）')

doc.add_heading('2.8 语法分析树', level=2)
doc.add_paragraph(
    '语法分析器在归约过程中同步构建了语法分析树。'
    '每个树节点包含符号名、是否为终结符、子节点列表和源码行号。'
    '最终的语法树以缩进文本的形式输出，使用 | 连线表示树结构。'
    '以下是一个简单赋值语句 "x = 10 ;" 的语法树示例：'
)

add_code_block(doc,
    '├─ P\n'
    '│  ├─ S\n'
    '│  │  ├─ id  (line 18)\n'
    '│  │  ├─ =   (line 18)\n'
    '│  │  ├─ E\n'
    '│  │  │  ├─ num  (line 18)\n'
    '│  ├─ ;  (line 18)'
)

doc.add_paragraph()
doc.add_paragraph(
    '（完整的语法分析树已输出至 output/Parser/parse_tree.txt 文件，'
    '包含测试用例中所有语句的完整语法树结构。）'
)

# ══════════════════════════════════════════════
# 三、编译器整体架构
# ══════════════════════════════════════════════
doc.add_heading('三、编译器整体架构', level=1)

doc.add_heading('3.1 模块间调用关系', level=2)
doc.add_paragraph(
    '经过两周的开发，编译器的整体架构已经清晰。各模块之间的调用关系如下图所示：'
)

doc.add_paragraph()
doc.add_paragraph('[架构图描述：编译器模块调用关系图]')
doc.add_paragraph()

add_code_block(doc,
    '┌─────────────────────────────────────────────────────────────────────┐\n'
    '│                          main.cpp                                  │\n'
    '│                    （程序入口，三阶段串联）                         │\n'
    '└────────┬────────────────────┬───────────────────┬──────────────────┘\n'
    '         │                    │                   │\n'
    '         ▼                    ▼                   ▼\n'
    '┌─────────────────┐  ┌────────────────┐  ┌───────────────────┐\n'
    '│  阶段一：        │  │  阶段二：       │  │  阶段三：          │\n'
    '│  词法分析        │  │  语法分析       │  │  目标代码生成      │\n'
    '│  (Lexer)        │  │  + 语义分析     │  │  (SemanticAnalyzer)│\n'
    '└────────┬────────┘  └───────┬────────┘  └─────────┬─────────┘\n'
    '         │                   │                     │\n'
    '         ▼                   ▼                     ▼\n'
    '┌─────────────────┐  ┌────────────────┐  ┌───────────────────┐\n'
    '│  ErrorHandler    │  │   Grammar      │  │  四元式 → Jack VM │\n'
    '│  SymbolTable     │  │  (文法定义)    │  │  (代码翻译)       │\n'
    '│  Token.h         │  │       │        │  │                   │\n'
    '└─────────────────┘  │       ▼        │  └───────────────────┘\n'
    '                     │  PrecedenceParser│\n'
    '                     │  (移进-归约)    │\n'
    '                     │       │         │\n'
    '                     │       ▼         │\n'
    '                     │  SemanticAnalyzer│\n'
    '                     │  (语义动作)     │\n'
    '                     └────────────────┘'
)

doc.add_paragraph()

doc.add_heading('3.2 数据流', level=2)
doc.add_paragraph('编译器的数据流方向如下：')

add_code_block(doc,
    '源代码 (.txt)\n'
    '      │\n'
    '      ▼\n'
    '┌───────────┐    ┌──────────────┐\n'
    '│   Lexer   │───→│ Token 序列   │\n'
    '└───────────┘    └──────┬───────┘\n'
    '                        │\n'
    '                        ▼\n'
    '               ┌────────────────────┐\n'
    '               │   Grammar          │\n'
    '               │  (优先关系表)      │\n'
    '               └────────┬───────────┘\n'
    '                        │\n'
    '                        ▼\n'
    '               ┌────────────────────┐\n'
    '               │  PrecedenceParser  │\n'
    '               │  (移进-归约分析)   │\n'
    '               └────────┬───────────┘\n'
    '                        │\n'
    '           ┌────────────┼────────────┐\n'
    '           ▼            ▼            ▼\n'
    '    ┌──────────┐  ┌──────────┐  ┌──────────────┐\n'
    '    │ 语法树   │  │ 跟踪日志 │  │ SemanticAnalyzer│\n'
    '    │ .txt     │  │ .txt     │  │ (四元式生成) │\n'
    '    └──────────┘  └──────────┘  └──────┬───────┘\n'
    '                                       │\n'
    '                                       ▼\n'
    '                               ┌──────────────┐\n'
    '                               │  Jack VM     │\n'
    '                               │  目标代码    │\n'
    '                               └──────────────┘'
)

# ══════════════════════════════════════════════
# 四、遇到的问题与解决方案
# ══════════════════════════════════════════════
doc.add_heading('四、遇到的问题与解决方案', level=1)

doc.add_heading('4.1 赋值语句与分号的优先关系冲突', level=2)
doc.add_paragraph(
    '问题描述：在算符优先分析中，赋值语句 "id = E ;" 的分析面临一个关键问题。'
    '等号（=）和分号（;）之间可能没有直接的优先关系（因为 E 可以推导出多种复杂的表达式），'
    '导致分析器在遇到分号时不知道应该归约 "id = E" 还是报告错误。'
)
doc.add_paragraph(
    '解决方案：在主控循环中加入了特殊的分号处理逻辑。当标准查表无法确定关系（关系为空）'
    '且当前输入为分号时，分析器会在栈中向下搜索 "id = E" 模式。'
    '如果找到匹配的模式，则强制执行归约操作，将 "id = E" 归约为 S。'
    '同时，在优先关系表中手动添加了 = > ; 的关系，确保大多数情况下的正常处理。'
    '这种"向下搜索 + 强制归约"的机制是对标准算符优先分析的有效扩展。'
)

doc.add_heading('4.2 结构体成员访问的多义性', level=2)
doc.add_paragraph(
    '问题描述：在 Micro-C 中，点号（.）同时用于结构体成员访问（如 p.x）和成员赋值（如 p.x = 10）。'
    '当分析器看到 "id . id" 时，需要判断这是成员访问表达式还是成员赋值语句的前缀。'
    '标准的算符优先分析无法通过简单的优先关系来区分这两种情况。'
)
doc.add_paragraph(
    '解决方案：在主控循环中加入了前瞻判断机制。当栈顶为 "id . id" 时，'
    '分析器会预读下一个 Token：如果下一个是等号（=），则判断为成员赋值，继续移进等号；'
    '否则判断为成员访问表达式，立即将 "id . id" 归约为 E。'
    '这种"预读-决策"模式类似于 LL(k) 分析中的 k=1 前瞻，但被嵌入到了移进-归约框架中。'
)

doc.add_heading('4.3 数组下标表达式的归约时机', level=2)
doc.add_paragraph(
    '问题描述：数组下标表达式 "id[E]" 的分析需要处理多个子问题：'
    '(1) 当看到 id 后跟 [ 时，不能将 id 归约为 E（否则会丢失数组名信息）；'
    '(2) 当看到 [ E 后跟 ] 时，需要移进 ] 而非归约 E；'
    '(3) 当看到 [ E ] 后，需要正确归约为 E 并生成数组访问四元式。'
    '这些问题的根源在于标准算符优先分析的归约时机与数组语法的特殊性不匹配。'
)
doc.add_paragraph(
    '解决方案：在主控循环中加入了三个特殊处理分支：'
    '(1) 当栈顶为 id 且输入为 [ 时，强制移进 [（而非归约 id → E）；'
    '(2) 当栈顶为 [ E 且输入为 ] 时，强制移进 ]；'
    '(3) 当栈顶为 [ E ] 时，根据上下文判断是数组元素赋值还是数组访问，'
    '分别执行不同的处理逻辑。这三个特殊处理分支确保了数组语法的正确分析。'
)

doc.add_heading('4.4 结构体变量声明的识别', level=2)
doc.add_paragraph(
    '问题描述："struct Point p" 是结构体变量声明，其语法为 "struct id id"。'
    '在算符优先分析中，当栈中有 "struct id" 且输入为 id 时，'
    '标准算法可能会尝试将 id 归约为 E（因为 E → id），导致分析失败。'
)
doc.add_paragraph(
    '解决方案：在归约分支中加入了特殊判断。当检测到栈顶为 "struct id" 且当前输入为 id 时，'
    '优先执行移进操作（而非归约），将第三个 id 压入栈中，形成完整的 "struct id id" 模式。'
    '随后在归约阶段，该模式被匹配为产生式 S → struct id id，执行结构体变量声明的语义动作。'
)

doc.add_heading('4.5 while 循环的四元式重排', level=2)
doc.add_paragraph(
    '问题描述：while 循环的四元式生成需要特殊的处理顺序。'
    '标准的归约顺序是先归约条件表达式 E，再归约循环体 P，最后归约整个 while 语句。'
    '但这会导致四元式的顺序为：条件计算 → label L0 → if_false → 循环体 → goto → label L1，'
    '而正确的顺序应该是：label L0 → 条件计算 → if_false → 循环体 → goto L0 → label L1。'
)
doc.add_paragraph(
    '解决方案：在 while 语句的语义动作中，采用了"提取-重排-恢复"的策略。'
    '首先，将条件表达式生成的最后一条四元式从四元式列表中提取出来（extractLastQuads）；'
    '然后，生成循环开始标签 label L0；'
    '接着，将提取的条件四元式恢复到列表中（restoreQuads）；'
    '最后，生成条件跳转和循环体的四元式。'
    '这种策略巧妙地解决了四元式顺序的问题，无需修改底层的四元式生成机制。'
)

# ══════════════════════════════════════════════
# 五、本周工作总结
# ══════════════════════════════════════════════
doc.add_heading('五、本周工作总结', level=1)

doc.add_paragraph('本周的工作取得了重要的阶段性成果，主要体现在以下几个方面：')

summary_items = [
    '完成了 Micro-C 文法的正式定义，设计了 28 条产生式，覆盖了语言的所有语法结构。',
    '实现了 FIRSTVT 和 LASTVT 集合的自动计算（不动点迭代算法），以及优先关系表的自动构建。',
    '成功实现了算符优先语法分析器，包括标准的移进-归约主控循环和 6 种特殊处理机制。',
    '实现了语法制导翻译，在归约时执行语义动作，包括符号表管理、类型检查和四元式生成。',
    '实现了控制流的标签回填技术，正确处理了 while 循环和 if-else 选择的跳转指令。',
    '实现了四元式到 Jack VM 指令集的代码翻译，完成了从源代码到目标代码的完整编译流程。',
    '编写了综合测试用例，验证了语法分析器、语义分析器和代码生成器的正确性。',
    '完善了项目的输出机制，生成了产生式列表、FIRSTVT/LASTVT 集合、优先关系表、移进-归约跟踪、语法分析树、四元式序列和目标代码等完整的编译产物。',
]
for item in summary_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    '总体而言，本周的工作进度符合预期计划，语法分析器的设计与实现已经完成，'
    '编译器已经具备了从源代码到目标代码的完整编译能力。'
    '这标志着编译器前端（词法分析 + 语法分析 + 语义分析）和代码生成阶段的基本完成。'
)

# ══════════════════════════════════════════════
# 六、下周工作计划
# ══════════════════════════════════════════════
doc.add_heading('六、下周工作计划', level=1)

doc.add_paragraph('下周的工作将进入编译器的优化和完善阶段。具体计划如下：')

doc.add_heading('6.1 错误处理机制的增强', level=2)
plan_items_1 = [
    '实现恐慌模式（Panic Mode）错误恢复：当遇到语法错误时，跳过输入直到找到同步符号（分号或右大括号），然后恢复分析。',
    '增强错误信息的精确度：在错误报告中包含期望的符号类型，帮助程序员快速定位和修复错误。',
    '实现错误计数和限制：设置最大错误数，超过阈值时终止分析，避免错误信息过多导致信息过载。',
]
for item in plan_items_1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 优化与性能改进', level=2)
plan_items_2 = [
    '优化优先关系表的存储：考虑使用更高效的数据结构（如哈希表）替代当前的嵌套 map。',
    '减少不必要的字符串拷贝：在栈操作和四元式生成中，使用引用和移动语义减少内存分配。',
    '优化语法树的构建：考虑使用智能指针管理树节点的内存。',
]
for item in plan_items_2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.3 文档与测试完善', level=2)
plan_items_3 = [
    '编写更全面的测试用例：覆盖更多的边界情况和错误情况。',
    '完善项目文档：更新设计文档，记录语法分析器的设计思路、算法细节和使用说明。',
    '编写用户手册：说明如何使用编译器、如何添加新的语言特性等。',
]
for item in plan_items_3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.4 课程设计报告准备', level=2)
plan_items_4 = [
    '整理课程设计报告的框架：包括需求分析、总体设计、详细设计、测试报告等章节。',
    '准备答辩材料：制作演示文稿，准备编译器的演示和讲解。',
]
for item in plan_items_4:
    doc.add_paragraph(item, style='List Bullet')

# ══════════════════════════════════════════════
# 七、附录
# ══════════════════════════════════════════════
doc.add_heading('七、附录', level=1)

doc.add_heading('附录A：终结符与非终结符完整列表', level=2)

doc.add_paragraph('非终结符（共 3 个）：')
nt_data = [
    ['P', '程序（Program）', '由若干分号分隔的语句构成'],
    ['S', '语句（Statement）', '包括声明、赋值、控制流等'],
    ['E', '表达式（Expression）', '包括算术、关系、成员访问等'],
]
add_table_with_header(doc, ['符号', '名称', '说明'], nt_data)
doc.add_paragraph()

doc.add_paragraph('终结符（共 25 个）：')
term_data = [
    ['关键字', 'int, char, double, struct, if, else, while, return'],
    ['函数名', 'id（包含 main, read, write）'],
    ['基本元素', 'id, num, char'],
    ['算术运算符', '+, -, *, /'],
    ['关系运算符', '<, >, =='],
    ['赋值运算符', '='],
    ['成员访问', '.'],
    ['界符', '(, ), {, }, ;, [, ]'],
    ['栈底标记', '$'],
]
add_table_with_header(doc, ['类别', '符号'], term_data)
doc.add_paragraph()

doc.add_heading('附录B：产生式编号对照表', level=2)
doc.add_paragraph('（已在正文 2.2.1 节中列出，此处不再重复）')

doc.add_heading('附录C：测试用例源代码', level=2)
doc.add_paragraph('本项目使用的综合测试用例如下：')

test_code = '''// Micro-C 综合测试 — 覆盖全部语法结构

// ===== 1. 变量声明 =====
int x ;
int y ;
int z ;
char c ;
double d ;

// ===== 2. 结构体定义 + 变量声明 =====
struct Point {
    int px ;
    int py ;
} ;
struct Point p ;

// ===== 3. 赋值语句 =====
x = 10 ;
y = 20 ;
c = 'A' ;
d = 3 ;
z = 0 ;

// ===== 4. 成员赋值 =====
p . px = 100 ;
p . py = 200 ;

// ===== 5. 算术表达式 =====
z = x + y ;
z = x - y ;
z = x * y ;
z = x / y ;
z = x + y * 2 ;
z = x * y + 1 ;

// ===== 6. 关系表达式 =====
z = x < y ;
z = x > y ;
z = x == y ;

// ===== 7. 成员访问（读取）=====
z = p . px ;
z = p . py ;

// ===== 8. 输入输出 =====
read(x) ;
write(z) ;
write(x) ;

// ===== 9. 循环 =====
while(x < 10) { x = x + 1 ; } ;

// ===== 10. 选择 =====
if(x > 5) { y = 1 ; } else { y = 2 ; } ;

// ===== 11. 返回 =====
return 0 ;'''

add_code_block(doc, test_code)

# ── 保存 ──
output_path = r'D:\c_cpp\ShiXvn_Compiler\第二周工作周报.docx'
doc.save(output_path)
print(f'周报已生成: {output_path}')
