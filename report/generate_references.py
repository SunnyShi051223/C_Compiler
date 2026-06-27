# -*- coding: utf-8 -*-
"""
生成参考文献 Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_reference(doc, ref_text):
    """添加参考文献条目"""
    p = doc.add_paragraph()
    run = p.add_run(ref_text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Pt(-24)  # 悬挂缩进
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    return p

def create_references():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ==================== 标题 ====================
    add_heading(doc, '参考文献', 1)

    doc.add_paragraph()

    # ==================== 教材类 ====================
    add_heading(doc, '一、教材类', 2)

    textbooks = [
        '[1] 陈火旺, 刘春林, 谭庆平等. 程序设计语言编译原理[M]. 第4版. 北京: 国防工业出版社, 2010.',
        '[2] Alfred V. Aho, Monica S. Lam, Ravi Sethi, Jeffrey D. Ullman. 编译原理[M]. 赵建华, 郑滔, 戴新宇译. 北京: 机械工业出版社, 2009.',
        '[3] Kenneth C. Louden. 编译原理及实践[M]. 冯博琴, 冯岚译. 北京: 机械工业出版社, 2000.',
        '[4] Andrew W. Appel. 现代编译原理[M]. 赵克佳, 黄春译. 北京: 人民邮电出版社, 2006.',
        '[5] 陈意云, 张昱. 编译原理[M]. 第3版. 北京: 高等教育出版社, 2014.',
        '[6] 张素琴, 吕映芝, 蒋维杜等. 编译原理[M]. 第2版. 北京: 清华大学出版社, 2005.',
    ]

    for ref in textbooks:
        add_reference(doc, ref)

    doc.add_paragraph()

    # ==================== nand2tetris相关 ====================
    add_heading(doc, '二、nand2tetris相关', 2)

    nand2tetris = [
        '[7] Noam Nisan, Shimon Schocken. 计算机系统要素: 从零开始构建现代计算机[M]. 周维, 宋磊译. 北京: 人民邮电出版社, 2010.',
        '[8] Noam Nisan, Shimon Schocken. The Elements of Computing Systems: Building a Modern Computer from First Principles[M]. Cambridge: MIT Press, 2005.',
        '[9] nand2tetris官方网站. https://www.nand2tetris.org/[EB/OL]. 2024.',
    ]

    for ref in nand2tetris:
        add_reference(doc, ref)

    doc.add_paragraph()

    # ==================== 编译器实现相关 ====================
    add_heading(doc, '三、编译器实现相关', 2)

    compiler_impl = [
        '[10] Jack W. Crenshaw. 编译器构造[M]. 李忠译. 北京: 电子工业出版社, 2005.',
        '[11] Dick Grune, Kees van Reeuwijk, Henri E. Bal等. 现代编译器实现: Java[M]. 第2版. 北京: 人民邮电出版社, 2010.',
        '[12] Steven S. Muchnick. 高级编译器设计与实现[M]. 赵克佳, 沈志宇译. 北京: 机械工业出版社, 2003.',
        '[13] Charles N. Fischer, Richard J. LeBlanc. 编译器构造: C语言[M]. 李健军译. 北京: 人民邮电出版社, 2005.',
    ]

    for ref in compiler_impl:
        add_reference(doc, ref)

    doc.add_paragraph()

    # ==================== 形式语言与自动机 ====================
    add_heading(doc, '四、形式语言与自动机', 2)

    formal_lang = [
        '[14] John E. Hopcroft, Rajeev Motwani, Jeffrey D. Ullman. 自动机理论、语言和计算导论[M]. 第3版. 孙家骕译. 北京: 机械工业出版社, 2008.',
        '[15] Michael Sipser. 计算理论导引[M]. 第3版. 唐常杰译. 北京: 机械工业出版社, 2015.',
        '[16] 陈有祺. 形式语言与自动机[M]. 天津: 南开大学出版社, 2003.',
    ]

    for ref in formal_lang:
        add_reference(doc, ref)

    doc.add_paragraph()

    # ==================== 数据结构与算法 ====================
    add_heading(doc, '五、数据结构与算法', 2)

    data_struct = [
        '[17] Thomas H. Cormen, Charles E. Leiserson, Ronald L. Rivest, Clifford Stein. 算法导论[M]. 第3版. 殷建平, 徐云, 王刚等译. 北京: 机械工业出版社, 2012.',
        '[18] 严蔚敏, 吴伟民. 数据结构(C语言版)[M]. 北京: 清华大学出版社, 2007.',
        '[19] Mark Allen Weiss. 数据结构与算法分析: C语言描述[M]. 第2版. 冯舜玺译. 北京: 机械工业出版社, 2004.',
    ]

    for ref in data_struct:
        add_reference(doc, ref)

    doc.add_paragraph()

    # ==================== C++编程 ====================
    add_heading(doc, '六、C++编程', 2)

    cpp_refs = [
        '[20] Bjarne Stroustrup. C++程序设计语言[M]. 第4版. 王刚, 杨巨峰译. 北京: 机械工业出版社, 2016.',
        '[21] Stanley B. Lippman, Josée Lajoie, Barbara E. Moo. C++ Primer[M]. 第5版. 王刚, 杨巨峰译. 北京: 电子工业出版社, 2013.',
        '[22] Scott Meyers. Effective C++[M]. 第3版. 侯捷译. 北京: 电子工业出版社, 2006.',
    ]

    for ref in cpp_refs:
        add_reference(doc, ref)

    doc.add_paragraph()

    # ==================== 软件工程 ====================
    add_heading(doc, '七、软件工程', 2)

    swe_refs = [
        '[23] Roger S. Pressman. 软件工程: 实践者的研究方法[M]. 第8版. 郑人杰, 马素霞译. 北京: 机械工业出版社, 2016.',
        '[24] Ian Sommerville. 软件工程[M]. 第10版. 程成译. 北京: 机械工业出版社, 2018.',
    ]

    for ref in swe_refs:
        add_reference(doc, ref)

    doc.add_paragraph()

    # ==================== 学术论文 ====================
    add_heading(doc, '八、学术论文', 2)

    papers = [
        '[25] Floyd R W. Syntactic analysis and operator precedence[J]. Journal of the ACM, 1963, 10(3): 316-333.',
        '[26] Dijkstra E W. A note on two problems in connexion with graphs[J]. Numerische Mathematik, 1959, 1(1): 269-271.',
        '[27] Aho A V, Johnson S C. LR parsing[J]. ACM Computing Surveys, 1974, 6(2): 99-124.',
        '[28] Knuth D E. On the translation of languages from left to right[J]. Information and Control, 1965, 8(6): 607-639.',
    ]

    for ref in papers:
        add_reference(doc, ref)

    doc.add_paragraph()

    # ==================== 网络资源 ====================
    add_heading(doc, '九、网络资源', 2)

    web_refs = [
        '[29] GCC官方文档. https://gcc.gnu.org/onlinedocs/[EB/OL]. 2024.',
        '[30] CMake官方文档. https://cmake.org/documentation/[EB/OL]. 2024.',
        '[31] C++参考手册. https://en.cppreference.com/[EB/OL]. 2024.',
        '[32] Git官方文档. https://git-scm.com/doc[EB/OL]. 2024.',
    ]

    for ref in web_refs:
        add_reference(doc, ref)

    doc.add_paragraph()
    doc.add_paragraph()

    # ==================== 参考文献格式说明 ====================
    add_heading(doc, '参考文献格式说明', 2)

    p = doc.add_paragraph()
    run = p.add_run('本报告参考文献采用GB/T 7714-2015《信息与文献 参考文献著录规则》格式，各类文献的著录格式如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 专著
    p = doc.add_paragraph()
    run = p.add_run('专著（图书）格式：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('[序号] 作者. 书名[M]. 版本. 出版地: 出版社, 出版年.')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.left_indent = Pt(36)

    doc.add_paragraph()

    # 译著
    p = doc.add_paragraph()
    run = p.add_run('译著格式：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('[序号] 原作者. 书名[M]. 译者译. 出版地: 出版社, 出版年.')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.left_indent = Pt(36)

    doc.add_paragraph()

    # 期刊论文
    p = doc.add_paragraph()
    run = p.add_run('期刊论文格式：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('[序号] 作者. 论文题目[J]. 期刊名, 年, 卷(期): 起止页码.')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.left_indent = Pt(36)

    doc.add_paragraph()

    # 电子资源
    p = doc.add_paragraph()
    run = p.add_run('电子资源格式：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('[序号] 作者. 题名[EB/OL]. 网址, 访问日期.')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.left_indent = Pt(36)

    doc.add_paragraph()

    # 文献类型标识
    p = doc.add_paragraph()
    run = p.add_run('文献类型标识说明：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    type_ids = [
        'M - 专著（Monograph）',
        'C - 论文集（Collection）',
        'N - 报纸文章（Newspaper）',
        'J - 期刊文章（Journal）',
        'D - 学位论文（Dissertation）',
        'R - 报告（Report）',
        'S - 标准（Standard）',
        'P - 专利（Patent）',
        'EB/OL - 电子资源（Electronic Bulletin Board / Online）',
    ]

    for type_id in type_ids:
        p = doc.add_paragraph()
        run = p.add_run(type_id)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(36)

    # 保存文档
    doc.save('D:\\c_cpp\\ShiXvn_Compiler\\report\\参考文献.docx')
    print('参考文献生成成功！')

if __name__ == '__main__':
    create_references()
