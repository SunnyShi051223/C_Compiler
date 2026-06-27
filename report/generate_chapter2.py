# -*- coding: utf-8 -*-
"""
生成第二章 系统分析 Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def add_heading(doc, text, level):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_paragraph(doc, text, bold_prefix=None):
    """添加段落"""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(text)
    else:
        run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_note(doc, text):
    """添加图表说明（红色斜体）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def create_chapter2():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ==================== 第二章标题 ====================
    add_heading(doc, '第二章  系统分析', 1)

    p = doc.add_paragraph()
    run = p.add_run('系统分析是编译器开发的重要阶段，主要包括源语言定义、目标语言定义和编译器功能需求分析。本章将详细定义mico C语言的词法规则和语法规则，以及Hack虚拟机的VM指令集，并分析编译器需要实现的功能。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 2.1 源语言定义 ====================
    add_heading(doc, '2.1 源语言定义', 2)

    p = doc.add_paragraph()
    run = p.add_run('本项目定义的源语言称为mico C（微型C语言），它是一种简化的C语言子集。mico C语言保留了C语言的核心语法特性，同时简化了部分复杂的语法特性，使其更适合作为编译器课程的实践教学内容。下面分别从词法规则和语法规则两个方面对mico C语言进行定义。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.1.1 词法规则
    add_heading(doc, '2.1.1 词法规则', 3)

    p = doc.add_paragraph()
    run = p.add_run('词法规则定义了源语言中单词（Token）的构成规则。mico C语言的单词分为以下几类：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 1. 关键字
    add_paragraph(doc, '关键字（Keyword）', bold_prefix='（1）')

    p = doc.add_paragraph()
    run = p.add_run('关键字是语言中预先定义的具有特殊含义的单词，不能用作标识符。mico C语言定义的关键字如表2-1所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-1 关键字列表
    p = doc.add_paragraph()
    run = p.add_run('表2-1 mico C语言关键字列表')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=12, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['序号', '关键字', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    keywords = [
        ('1', 'int', '整数类型关键字'),
        ('2', 'char', '字符类型关键字'),
        ('3', 'double', '双精度浮点类型关键字'),
        ('4', 'struct', '结构体定义关键字'),
        ('5', 'if', '条件语句关键字'),
        ('6', 'else', '条件语句分支关键字'),
        ('7', 'while', '循环语句关键字'),
        ('8', 'return', '函数返回关键字'),
        ('9', 'read', '输入语句关键字'),
        ('10', 'write', '输出语句关键字'),
        ('11', 'main', '主函数关键字'),
    ]

    for i, (col1, col2, col3) in enumerate(keywords, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 2. 标识符
    add_paragraph(doc, '标识符（Identifier）', bold_prefix='（2）')

    p = doc.add_paragraph()
    run = p.add_run('标识符用于命名变量、函数、结构体等程序实体。mico C语言的标识符规则如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    id_rules = [
        '• 标识符由字母、数字和下划线组成',
        '• 标识符必须以字母或下划线开头',
        '• 标识符不能是关键字',
        '• 标识符区分大小写',
    ]

    for rule in id_rules:
        p = doc.add_paragraph()
        run = p.add_run(rule)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(36)

    p = doc.add_paragraph()
    run = p.add_run('标识符的正则表达式定义为：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('identifier → letter (letter | digit)*')
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run('其中，letter → a | b | ... | z | A | B | ... | Z | _')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('      digit → 0 | 1 | 2 | ... | 9')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 3. 常量
    add_paragraph(doc, '常量（Constant）', bold_prefix='（3）')

    p = doc.add_paragraph()
    run = p.add_run('mico C语言支持以下两种常量：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('整型常量：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = p.add_run('由数字序列组成，表示整数值。例如：0、42、1024、-5等。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('整型常量的正则表达式定义为：number → digit+')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('字符常量：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = p.add_run('用单引号括起来的单个字符。例如：\'A\'、\'b\'、\'0\'、\'\\n\'等。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('字符常量的正则表达式定义为：char_literal → \' (letter | digit | escape) \'')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 4. 运算符
    add_paragraph(doc, '运算符（Operator）', bold_prefix='（4）')

    p = doc.add_paragraph()
    run = p.add_run('mico C语言支持的运算符如表2-2所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-2 运算符列表
    p = doc.add_paragraph()
    run = p.add_run('表2-2 mico C语言运算符列表')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=14, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['类别', '运算符', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    operators = [
        ('算术运算符', '+', '加法'),
        ('', '-', '减法'),
        ('', '*', '乘法'),
        ('', '/', '除法'),
        ('', '%', '取模'),
        ('关系运算符', '>', '大于'),
        ('', '<', '小于'),
        ('', '>=', '大于等于'),
        ('', '<=', '小于等于'),
        ('', '==', '等于'),
        ('', '!=', '不等于'),
        ('赋值运算符', '=', '赋值'),
        ('成员访问', '.', '结构体成员访问'),
    ]

    for i, (col1, col2, col3) in enumerate(operators, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 5. 分隔符
    add_paragraph(doc, '分隔符（Delimiter）', bold_prefix='（5）')

    p = doc.add_paragraph()
    run = p.add_run('mico C语言支持的分隔符如表2-3所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-3 分隔符列表
    p = doc.add_paragraph()
    run = p.add_run('表2-3 mico C语言分隔符列表')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['分隔符', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    delimiters = [
        ('(', '左括号'),
        (')', '右括号'),
        ('{', '左花括号'),
        ('}', '右花括号'),
        ('[', '左方括号'),
        (']', '右方括号'),
        (';', '分号'),
    ]

    for i, (col1, col2) in enumerate(delimiters, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        for j in range(2):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 6. 空白字符和注释
    add_paragraph(doc, '空白字符和注释', bold_prefix='（6）')

    p = doc.add_paragraph()
    run = p.add_run('空白字符包括空格、制表符、换行符等，在词法分析阶段被忽略（用于分隔单词）。mico C语言不支持注释。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Token类型汇总
    p = doc.add_paragraph()
    run = p.add_run('mico C语言的Token类型汇总如表2-4所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-4 Token类型汇总
    p = doc.add_paragraph()
    run = p.add_run('表2-4 Token类型汇总')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Token类型', '类型标识', '示例']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    token_types = [
        ('关键字', 'KEYWORD', 'int, char, if, while'),
        ('标识符', 'IDENTIFIER', 'x, myVar, func1'),
        ('整型常数', 'NUMBER', '0, 42, 1024'),
        ('字符常数', 'CHAR_LITERAL', '\'A\', \'\\n\''),
        ('运算符', 'OPERATOR', '+, -, *, =, =='),
        ('分隔符', 'DELIMITER', '(, ), {, }, ;'),
        ('文件结束', 'END_OF_FILE', '(EOF)'),
    ]

    for i, (col1, col2, col3) in enumerate(token_types, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 2.1.2 语法规则
    add_heading(doc, '2.1.2 语法规则', 3)

    p = doc.add_paragraph()
    run = p.add_run('语法规则定义了源语言中程序的结构和组成方式。mico C语言的语法规则采用BNF范式（Backus-Naur Form）进行定义。下面分别介绍程序结构、声明语句、赋值语句、控制结构、表达式等各个部分的语法规则。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 1. 程序结构
    add_paragraph(doc, '程序结构', bold_prefix='（1）')

    p = doc.add_paragraph()
    run = p.add_run('mico C语言的程序由一系列语句组成，语句之间用分号分隔。程序的起始符号为P（Program）。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('程序结构的产生式定义如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    productions = [
        'P → S ; P      （多个语句）',
        'P → S ;        （单个语句）',
    ]

    for prod in productions:
        p = doc.add_paragraph()
        run = p.add_run(prod)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    # 2. 声明语句
    add_paragraph(doc, '声明语句', bold_prefix='（2）')

    p = doc.add_paragraph()
    run = p.add_run('声明语句用于声明变量、数组和结构体。mico C语言支持以下声明语句：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('变量声明：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    productions = [
        'S → int id       （声明整型变量）',
        'S → char id      （声明字符变量）',
        'S → double id    （声明双精度浮点变量）',
    ]

    for prod in productions:
        p = doc.add_paragraph()
        run = p.add_run(prod)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    p = doc.add_paragraph()
    run = p.add_run('数组声明：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    productions = [
        'S → int id [ num ]     （声明整型数组）',
        'S → char id [ num ]    （声明字符数组）',
        'S → double id [ num ]  （声明双精度浮点数组）',
    ]

    for prod in productions:
        p = doc.add_paragraph()
        run = p.add_run(prod)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    p = doc.add_paragraph()
    run = p.add_run('结构体声明：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    productions = [
        'S → struct id { P }     （定义结构体类型）',
        'S → struct id id        （声明结构体变量）',
    ]

    for prod in productions:
        p = doc.add_paragraph()
        run = p.add_run(prod)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    # 3. 赋值语句
    add_paragraph(doc, '赋值语句', bold_prefix='（3）')

    p = doc.add_paragraph()
    run = p.add_run('赋值语句用于将表达式的值赋给变量。赋值语句的产生式定义如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('S → id = E      （将表达式E的值赋给变量id）')
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.left_indent = Pt(72)

    # 4. 控制结构
    add_paragraph(doc, '控制结构', bold_prefix='（4）')

    p = doc.add_paragraph()
    run = p.add_run('mico C语言支持if-else条件语句和while循环语句。控制结构的产生式定义如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('if-else条件语句：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    productions = [
        'S → if ( E ) { P } else { P }   （带else分支的条件语句）',
        'S → if ( E ) { P }              （不带else分支的条件语句）',
    ]

    for prod in productions:
        p = doc.add_paragraph()
        run = p.add_run(prod)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    p = doc.add_paragraph()
    run = p.add_run('while循环语句：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('S → while ( E ) { P }   （while循环语句）')
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.left_indent = Pt(72)

    p = doc.add_paragraph()
    run = p.add_run('return语句：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('S → return E   （函数返回语句）')
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.left_indent = Pt(72)

    # 5. 表达式
    add_paragraph(doc, '表达式', bold_prefix='（5）')

    p = doc.add_paragraph()
    run = p.add_run('表达式是计算值的基本单位。mico C语言支持以下类型的表达式：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('函数调用表达式：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    productions = [
        'E → id ( )              （无参函数调用）',
        'E → id ( E )            （单参函数调用）',
        'E → id ( E , E )        （双参函数调用）',
        'E → id ( E , E , E )    （三参函数调用）',
    ]

    for prod in productions:
        p = doc.add_paragraph()
        run = p.add_run(prod)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    p = doc.add_paragraph()
    run = p.add_run('算术和关系表达式：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    productions = [
        'E → E + E      （加法）',
        'E → E - E      （减法）',
        'E → E * E      （乘法）',
        'E → E / E      （除法）',
        'E → E < E      （小于比较）',
        'E → E > E      （大于比较）',
    ]

    for prod in productions:
        p = doc.add_paragraph()
        run = p.add_run(prod)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    p = doc.add_paragraph()
    run = p.add_run('基本表达式：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    productions = [
        'E → ( E )      （括号表达式）',
        'E → id         （变量引用）',
        'E → num        （整型常量）',
        'E → char       （字符常量）',
    ]

    for prod in productions:
        p = doc.add_paragraph()
        run = p.add_run(prod)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    # 6. 产生式汇总
    add_paragraph(doc, '产生式汇总', bold_prefix='（6）')

    p = doc.add_paragraph()
    run = p.add_run('mico C语言的完整产生式列表如表2-5所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-5 产生式汇总
    p = doc.add_paragraph()
    run = p.add_run('表2-5 mico C语言产生式汇总')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=28, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['编号', '产生式', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    productions_data = [
        ('1', 'P → S ; P', '程序（多个语句）'),
        ('2', 'P → S ;', '程序（单个语句）'),
        ('3', 'S → int id', '声明整型变量'),
        ('4', 'S → char id', '声明字符变量'),
        ('5', 'S → double id', '声明双精度浮点变量'),
        ('6', 'S → int id [ num ]', '声明整型数组'),
        ('7', 'S → char id [ num ]', '声明字符数组'),
        ('8', 'S → double id [ num ]', '声明双精度浮点数组'),
        ('9', 'S → struct id { P }', '定义结构体类型'),
        ('10', 'S → struct id id', '声明结构体变量'),
        ('11', 'S → id = E', '赋值语句'),
        ('12', 'S → E', '表达式语句'),
        ('13', 'S → while ( E ) { P }', 'while循环'),
        ('14', 'S → if ( E ) { P } else { P }', 'if-else语句'),
        ('15', 'S → if ( E ) { P }', 'if语句'),
        ('16', 'S → return E', 'return语句'),
        ('17', 'E → id ( )', '无参函数调用'),
        ('18', 'E → id ( E )', '单参函数调用'),
        ('19', 'E → id ( E , E )', '双参函数调用'),
        ('20', 'E → id ( E , E , E )', '三参函数调用'),
        ('21', 'E → E + E', '加法'),
        ('22', 'E → E - E', '减法'),
        ('23', 'E → E * E', '乘法'),
        ('24', 'E → E / E', '除法'),
        ('25', 'E → E < E', '小于比较'),
        ('26', 'E → E > E', '大于比较'),
        ('27', 'E → ( E )', '括号表达式'),
    ]

    for i, (col1, col2, col3) in enumerate(productions_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                if j == 1:
                    for run in paragraph.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 7. 文法符号说明
    add_paragraph(doc, '文法符号说明', bold_prefix='（7）')

    p = doc.add_paragraph()
    run = p.add_run('mico C语言文法中的符号说明如表2-6所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-6 文法符号说明
    p = doc.add_paragraph()
    run = p.add_run('表2-6 文法符号说明')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['符号', '类型', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    symbols_data = [
        ('P', '非终结符', '程序（Program）'),
        ('S', '非终结符', '语句（Statement）'),
        ('E', '非终结符', '表达式（Expression）'),
        ('id', '终结符', '标识符（Identifier）'),
        ('num', '终结符', '整型常量（Number）'),
    ]

    for i, (col1, col2, col3) in enumerate(symbols_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 8. 简单优先文法说明
    add_paragraph(doc, '简单优先文法说明', bold_prefix='（8）')

    p = doc.add_paragraph()
    run = p.add_run('本项目采用简单优先法进行语法分析，因此文法需要满足简单优先文法的条件：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    conditions = [
        '• 右部互不相同：任何两个产生式的右部不能完全相同',
        '• 无ε产生式：产生式的右部不能为空',
        '• 无相邻非终结符：产生式右部中不能有两个非终结符相邻',
    ]

    for cond in conditions:
        p = doc.add_paragraph()
        run = p.add_run(cond)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(36)

    p = doc.add_paragraph()
    run = p.add_run('本项目的文法设计满足以上三个条件，可以使用简单优先法进行语法分析。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 2.2 目标语言定义 ====================
    add_heading(doc, '2.2 目标语言定义', 2)

    p = doc.add_paragraph()
    run = p.add_run('本项目选择Hack虚拟机作为目标平台，使用VM指令集作为目标语言。Hack虚拟机是一种基于栈的虚拟机，由nand2tetris课程提出，它提供了一个简洁而完整的虚拟机架构。下面详细介绍Hack虚拟机的架构和VM指令集。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.2.1 Hack虚拟机架构
    add_heading(doc, '2.2.1 Hack虚拟机架构', 3)

    p = doc.add_paragraph()
    run = p.add_run('Hack虚拟机是一种基于栈的虚拟机，其主要特点包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    features = [
        '基于栈的架构：Hack虚拟机使用栈来存储操作数和中间结果。大多数指令都是从栈中取操作数，将结果压入栈中。栈指针（SP）存储在RAM[0]中，指向栈顶元素的下一个位置。',
        '内存分段：Hack虚拟机将内存分为以下几个段：',
        '  - 栈段（Stack）：存储函数的局部变量、参数和中间计算结果',
        '  - 堆段（Heap）：存储动态分配的内存（如数组、结构体）',
        '  - 静态段（Static）：存储静态变量',
        '  - 常量段（Constant）：虚拟段，用于表示常量值',
        '  - 指针段（Pointer）：存储this和that寄存器的值',
        '  - 临时段（Temp）：存储临时变量（R5-R12）',
        '函数调用机制：Hack虚拟机支持函数的定义和调用，使用栈帧来管理函数的局部变量和参数。',
    ]

    for feature in features:
        p = doc.add_paragraph()
        run = p.add_run(feature)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if feature.startswith('  '):
            p.paragraph_format.left_indent = Pt(72)
        else:
            p.paragraph_format.left_indent = Pt(36)

    p = doc.add_paragraph()
    run = p.add_run('Hack虚拟机的内存布局如图2-1所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_note(doc, '【图表说明】请绘制一个Hack虚拟机内存布局图，展示以下内容：RAM[0]存储栈指针SP，RAM[1-2]存储LCL和ARG寄存器，RAM[3-4]存储THIS和THAT寄存器，RAM[5-12]为临时段，RAM[13-15]为保留段，RAM[16-255]为静态段，RAM[256-2047]为栈段，RAM[2048-16383]为堆段。')

    doc.add_paragraph()

    # 2.2.2 VM指令集
    add_heading(doc, '2.2.2 VM指令集', 3)

    p = doc.add_paragraph()
    run = p.add_run('Hack虚拟机的VM指令集包含以下几类指令：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 1. 栈操作指令
    add_paragraph(doc, '栈操作指令', bold_prefix='（1）')

    p = doc.add_paragraph()
    run = p.add_run('栈操作指令用于在栈和内存段之间传输数据。栈操作指令如表2-7所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-7 栈操作指令
    p = doc.add_paragraph()
    run = p.add_run('表2-7 栈操作指令')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['指令', '格式', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    stack_ops = [
        ('push', 'push segment index', '将指定段的第index个元素压入栈顶'),
        ('pop', 'pop segment index', '将栈顶元素弹出到指定段的第index个位置'),
    ]

    for i, (col1, col2, col3) in enumerate(stack_ops, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('其中，segment可以是以下内存段之一：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 内存段说明
    segments = [
        'local：局部变量段，映射到LCL寄存器指向的内存区域',
        'argument：参数段，映射到ARG寄存器指向的内存区域',
        'this：this段，映射到THIS寄存器指向的内存区域',
        'that：that段，映射到THAT寄存器指向的内存区域',
        'constant：常量段（虚拟段），push constant x将常量x压入栈',
        'static：静态段，存储静态变量',
        'temp：临时段，映射到RAM[5-12]',
        'pointer：指针段，pointer 0访问THIS，pointer 1访问THAT',
    ]

    for seg in segments:
        p = doc.add_paragraph()
        run = p.add_run(seg)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(36)

    # 2. 算术逻辑指令
    add_paragraph(doc, '算术逻辑指令', bold_prefix='（2）')

    p = doc.add_paragraph()
    run = p.add_run('算术逻辑指令用于执行算术运算和逻辑运算。算术逻辑指令如表2-8所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-8 算术逻辑指令
    p = doc.add_paragraph()
    run = p.add_run('表2-8 算术逻辑指令')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['指令', '操作数', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    alu_ops = [
        ('add', '2个', 'x + y'),
        ('sub', '2个', 'x - y'),
        ('neg', '1个', '-y'),
        ('eq', '2个', 'x == y（返回-1表示真，0表示假）'),
        ('gt', '2个', 'x > y'),
        ('lt', '2个', 'x < y'),
        ('and', '2个', 'x AND y（按位与）'),
        ('or', '2个', 'x OR y（按位或）'),
        ('not', '1个', 'NOT y（按位非）'),
    ]

    for i, (col1, col2, col3) in enumerate(alu_ops, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 3. 流程控制指令
    add_paragraph(doc, '流程控制指令', bold_prefix='（3）')

    p = doc.add_paragraph()
    run = p.add_run('流程控制指令用于控制程序的执行流程。流程控制指令如表2-9所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-9 流程控制指令
    p = doc.add_paragraph()
    run = p.add_run('表2-9 流程控制指令')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['指令', '格式', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    flow_ops = [
        ('label', 'label labelName', '定义标签labelName'),
        ('goto', 'goto labelName', '无条件跳转到标签labelName'),
        ('if-goto', 'if-goto labelName', '如果栈顶值非0，跳转到标签labelName'),
    ]

    for i, (col1, col2, col3) in enumerate(flow_ops, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 4. 函数调用指令
    add_paragraph(doc, '函数调用指令', bold_prefix='（4）')

    p = doc.add_paragraph()
    run = p.add_run('函数调用指令用于定义和调用函数。函数调用指令如表2-10所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-10 函数调用指令
    p = doc.add_paragraph()
    run = p.add_run('表2-10 函数调用指令')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['指令', '格式', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    func_ops = [
        ('function', 'function f n', '定义函数f，包含n个局部变量'),
        ('call', 'call f m', '调用函数f，传递m个参数'),
        ('return', 'return', '从当前函数返回，返回值在栈顶'),
    ]

    for i, (col1, col2, col3) in enumerate(func_ops, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('函数调用的栈帧结构如图2-2所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_note(doc, '【图表说明】请绘制一个函数调用栈帧结构图，展示以下内容：调用函数时，需要将返回地址、调用者的LCL、ARG、THIS、THAT压入栈中，然后设置新的LCL和ARG指针，最后为局部变量分配空间。')

    doc.add_paragraph()

    # 5. VM指令汇总
    add_paragraph(doc, 'VM指令汇总', bold_prefix='（5）')

    p = doc.add_paragraph()
    run = p.add_run('Hack虚拟机的完整VM指令集如表2-11所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-11 VM指令汇总
    p = doc.add_paragraph()
    run = p.add_run('表2-11 VM指令汇总')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=20, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['类别', '指令', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    vm_instructions = [
        ('栈操作', 'push segment index', '压栈'),
        ('', 'pop segment index', '出栈'),
        ('算术逻辑', 'add', '加法'),
        ('', 'sub', '减法'),
        ('', 'neg', '取反'),
        ('', 'eq', '等于比较'),
        ('', 'gt', '大于比较'),
        ('', 'lt', '小于比较'),
        ('', 'and', '逻辑与'),
        ('', 'or', '逻辑或'),
        ('', 'not', '逻辑非'),
        ('流程控制', 'label labelName', '定义标签'),
        ('', 'goto labelName', '无条件跳转'),
        ('', 'if-goto labelName', '条件跳转'),
        ('函数', 'function f n', '定义函数'),
        ('', 'call f m', '调用函数'),
        ('', 'return', '返回'),
        ('内存', 'push constant x', '压入常量'),
        ('', 'pop static i', '弹出到静态变量'),
    ]

    for i, (col1, col2, col3) in enumerate(vm_instructions, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                if j == 1:
                    for run in paragraph.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 2.2.3 目标代码示例
    add_heading(doc, '2.2.3 目标代码示例', 3)

    p = doc.add_paragraph()
    run = p.add_run('下面通过一个简单的例子来说明源代码到目标代码的转换过程。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('源代码：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    source_code = [
        'int a = 10;',
        'int b = 20;',
        'int c = a + b;',
        'write(c);',
    ]

    for line in source_code:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    p = doc.add_paragraph()
    run = p.add_run('生成的VM代码：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    vm_code = [
        'function Main.main 3',
        'push constant 10',
        'pop local 0',
        'push constant 20',
        'pop local 1',
        'push local 0',
        'push local 1',
        'add',
        'pop local 2',
        'push local 2',
        'call Output.printInt 1',
        'pop temp 0',
    ]

    for line in vm_code:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)

    p = doc.add_paragraph()
    run = p.add_run('代码说明：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    explanations = [
        '• function Main.main 3：定义Main.main函数，包含3个局部变量（a, b, c）',
        '• push constant 10 / pop local 0：将常量10赋值给局部变量a（local 0）',
        '• push constant 20 / pop local 1：将常量20赋值给局部变量b（local 1）',
        '• push local 0 / push local 1 / add：将a和b的值压入栈并执行加法',
        '• pop local 2：将加法结果赋值给局部变量c（local 2）',
        '• push local 2 / call Output.printInt 1：调用Output.printInt函数输出c的值',
        '• pop temp 0：丢弃函数返回值',
    ]

    for exp in explanations:
        p = doc.add_paragraph()
        run = p.add_run(exp)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(36)

    # ==================== 2.3 编译器功能需求 ====================
    add_heading(doc, '2.3 编译器功能需求', 2)

    p = doc.add_paragraph()
    run = p.add_run('根据实训任务书的要求和mico C语言的定义，编译器需要实现以下功能：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.3.1 词法分析功能
    add_heading(doc, '2.3.1 词法分析功能', 3)

    p = doc.add_paragraph()
    run = p.add_run('词法分析是编译器的第一个阶段，其主要功能包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    lexer_functions = [
        '（1）单词截取：从源程序的字符序列中识别出一个个单词（Token），单词的边界由分隔符（如空格、换行符、运算符等）确定。',
        '（2）词法错误识别：在单词截取过程中，识别并报告词法错误。常见的词法错误包括：非法字符、数字格式错误、字符常量未闭合等。',
        '（3）Token序列生成：将识别出的单词转换成Token序列，每个Token包含类型和值两个属性。Token类型包括：关键字（KEYWORD）、标识符（IDENTIFIER）、数字（NUMBER）、字符（CHAR_LITERAL）、运算符（OPERATOR）、分隔符（DELIMITER）等。',
        '（4）符号表初始化：在词法分析过程中，将标识符信息添加到符号表中，为后续的语义分析做准备。',
        '（5）为语法分析做准备：将生成的Token序列传递给语法分析器，作为语法分析的输入。',
    ]

    for func in lexer_functions:
        p = doc.add_paragraph()
        run = p.add_run(func)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.3.2 语法分析功能
    add_heading(doc, '2.3.2 语法分析功能', 3)

    p = doc.add_paragraph()
    run = p.add_run('语法分析是编译器的第二个阶段，其主要功能包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    parser_functions = [
        '（1）文法处理：定义mico C语言的文法，包括产生式、终结符、非终结符等。计算FIRSTVT和LASTVT集合，构建优先关系表。',
        '（2）语法分析：使用简单优先法对Token序列进行语法分析，判断源程序是否符合语法规则。简单优先法通过比较栈顶符号和当前输入符号的优先关系来决定移进或归约操作。',
        '（3）语法树构建：在语法分析过程中，构建语法树（Parse Tree）和抽象语法树（AST），为后续的语义分析和代码生成提供基础。',
        '（4）错误处理：在语法分析过程中，识别并报告语法错误，具备容错功能。当遇到语法错误时，能够给出错误位置和错误类型信息。',
        '（5）分析过程记录：记录语法分析的详细过程，包括移进、归约、接受等操作，便于调试和理解。',
    ]

    for func in parser_functions:
        p = doc.add_paragraph()
        run = p.add_run(func)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.3.3 语义分析功能
    add_heading(doc, '2.3.3 语义分析功能', 3)

    p = doc.add_paragraph()
    run = p.add_run('语义分析是编译器的第三个阶段，其主要功能包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    semantic_functions = [
        '（1）类型检查：检查表达式的类型是否正确，包括算术运算的类型兼容性、赋值语句的类型匹配、函数参数的类型检查等。',
        '（2）符号表管理：维护符号表，记录变量、函数等标识符的信息，包括类型、作用域、地址等。支持变量的声明、查找和作用域管理。',
        '（3）语义动作执行：在语法分析过程中，当产生式被归约时，执行相应的语义动作。语义动作可以进行语义检查、生成中间代码等操作。',
        '（4）中间代码生成：生成四元式（Quadruple）作为中间代码。四元式是一种三地址码，每个四元式包含四个部分：运算符、操作数1、操作数2、结果。',
        '（5）错误处理：识别并报告语义错误，如未声明的变量、类型不匹配、函数参数错误等。',
    ]

    for func in semantic_functions:
        p = doc.add_paragraph()
        run = p.add_run(func)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.3.4 代码生成功能
    add_heading(doc, '2.3.4 代码生成功能', 3)

    p = doc.add_paragraph()
    run = p.add_run('代码生成是编译器的最后一个阶段，其主要功能包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    codegen_functions = [
        '（1）指令选择：将四元式转换成对应的VM指令。不同的四元式对应不同的VM指令序列。',
        '（2）内存管理：管理局部变量、参数、临时变量的内存分配。局部变量分配在local段，参数分配在argument段。',
        '（3）表达式代码生成：生成算术运算、比较运算、逻辑运算的VM指令。运算指令从栈中取操作数，将结果压入栈中。',
        '（4）控制流代码生成：生成if-else、while等控制结构的VM指令。使用标签和跳转指令实现控制流。',
        '（5）函数调用代码生成：生成函数定义、函数调用、函数返回的VM指令。处理参数传递、栈帧管理等。',
        '（6）输入输出代码生成：生成read、write语句的VM指令。调用操作系统的输入输出函数。',
    ]

    for func in codegen_functions:
        p = doc.add_paragraph()
        run = p.add_run(func)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.3.5 功能需求汇总
    add_heading(doc, '2.3.5 功能需求汇总', 3)

    p = doc.add_paragraph()
    run = p.add_run('编译器的功能需求汇总如表2-12所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表2-12 功能需求汇总
    p = doc.add_paragraph()
    run = p.add_run('表2-12 编译器功能需求汇总')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=17, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['阶段', '功能', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    requirements_data = [
        ('词法分析', '单词截取', '从字符序列中识别Token'),
        ('', '词法错误识别', '识别非法字符、格式错误等'),
        ('', 'Token序列生成', '生成Token序列供语法分析使用'),
        ('', '符号表初始化', '将标识符添加到符号表'),
        ('语法分析', '文法处理', '定义文法、计算FIRSTVT/LASTVT'),
        ('', '语法分析', '使用简单优先法分析语法'),
        ('', '语法树构建', '构建Parse Tree和AST'),
        ('', '错误处理', '识别和报告语法错误'),
        ('语义分析', '类型检查', '检查类型兼容性'),
        ('', '符号表管理', '管理变量、函数信息'),
        ('', '中间代码生成', '生成四元式'),
        ('代码生成', '指令选择', '四元式转换为VM指令'),
        ('', '内存管理', '管理局部变量、参数内存'),
        ('', '表达式代码生成', '生成算术、逻辑运算指令'),
        ('', '控制流代码生成', '生成if-else、while指令'),
        ('', '函数调用代码生成', '生成函数相关指令'),
    ]

    for i, (col1, col2, col3) in enumerate(requirements_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                if j == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # ==================== 2.4 本章小结 ====================
    add_heading(doc, '2.4 本章小结', 2)

    p = doc.add_paragraph()
    run = p.add_run('本章对mico C语言编译器进行了系统分析，主要包括源语言定义、目标语言定义和编译器功能需求分析三个方面。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('在源语言定义部分，详细定义了mico C语言的词法规则和语法规则。词法规则定义了关键字、标识符、常量、运算符、分隔符等单词的构成规则；语法规则采用BNF范式定义了程序结构、声明语句、赋值语句、控制结构、表达式等各个部分的产生式，并列出了完整的产生式汇总表。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('在目标语言定义部分，介绍了Hack虚拟机的架构和VM指令集。Hack虚拟机是一种基于栈的虚拟机，其指令集包括栈操作指令、算术逻辑指令、流程控制指令、函数调用指令等。通过一个简单的例子说明了源代码到目标代码的转换过程。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('在编译器功能需求分析部分，详细分析了词法分析、语法分析、语义分析、代码生成四个阶段需要实现的功能，并给出了功能需求汇总表。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('通过本章的系统分析，明确了编译器需要处理的语言特性和需要实现的功能，为后续章节的详细设计和实现奠定了基础。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存文档
    doc.save('D:\\c_cpp\\ShiXvn_Compiler\\report\\第二章_系统分析.docx')
    print('第二章生成成功！')

if __name__ == '__main__':
    create_chapter2()
