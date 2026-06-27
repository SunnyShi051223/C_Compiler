# -*- coding: utf-8 -*-
"""
生成第四章 实训小结 Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def add_heading(doc, text, level):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_paragraph(doc, text, bold_prefix=None, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(text)
    else:
        run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.left_indent = Pt(36)
    return p

def add_note(doc, text):
    """添加图表说明（红色斜体）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def create_chapter4():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ==================== 第四章标题 ====================
    add_heading(doc, '第四章  实训小结', 1)

    p = doc.add_paragraph()
    run = p.add_run('本章对基于简单优先法的微型程序语言编译器设计与实现的实训进行总结，包括设计特点、实训体会、改进建议和总结四个方面。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 4.1 设计特点 ====================
    add_heading(doc, '4.1 设计特点', 2)

    p = doc.add_paragraph()
    run = p.add_run('本编译器在设计和实现过程中，具有以下几个显著特点：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 4.1.1 简单优先法的应用
    add_heading(doc, '4.1.1 简单优先法的应用', 3)

    p = doc.add_paragraph()
    run = p.add_run('本编译器采用简单优先法作为语法分析方法，这是本项目最显著的设计特点之一。简单优先法是一种经典的自底向上语法分析方法，具有以下优点：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    advantages = [
        '（1）原理直观：简单优先法通过比较文法符号之间的优先关系来决定移进或归约操作，原理简单易懂，便于学习和理解。',
        '（2）实现简单：相比于LR分析法，简单优先法的实现更加简单，不需要构造复杂的项目集和状态机，降低了编译器开发的难度。',
        '（3）适用性广：简单优先法适用于大多数上下文无关文法，能够处理常见的程序语言语法结构。',
        '（4）易于调试：简单优先法的分析过程清晰明了，便于调试和错误定位。',
    ]

    for advantage in advantages:
        add_paragraph(doc, advantage, indent=True)

    p = doc.add_paragraph()
    run = p.add_run('在实现过程中，我们严格按照简单优先法的理论进行设计，包括计算FIRSTVT和LASTVT集合、构建优先关系表、实现移进-归约分析算法等。通过实践，我们深入理解了简单优先法的工作原理和实现细节。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 简单优先法与其他分析法的比较
    p = doc.add_paragraph()
    run = p.add_run('简单优先法与其他语法分析方法的比较如表4-1所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('表4-1 语法分析方法比较')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['分析方法', '分析方向', '文法类', '实现复杂度', '适用场景']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    comparison_data = [
        ('简单优先法', '自底向上', '简单优先文法', '中等', '教学、小型编译器'),
        ('LL分析法', '自顶向下', 'LL文法', '简单', '递归下降分析'),
        ('LR分析法', '自底向上', 'LR文法', '复杂', '大型编译器'),
        ('算符优先法', '自底向上', '算符优先文法', '中等', '表达式分析'),
    ]

    for i, (col1, col2, col3, col4, col5) in enumerate(comparison_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        table.rows[i].cells[3].text = col4
        table.rows[i].cells[4].text = col5
        for j in range(5):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 4.1.2 模块化设计
    add_heading(doc, '4.1.2 模块化设计', 3)

    p = doc.add_paragraph()
    run = p.add_run('本编译器采用模块化设计思想，将编译器划分为多个独立的模块，每个模块负责特定的功能。模块化设计具有以下优点：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    modularity_advantages = [
        '（1）职责清晰：每个模块都有明确的职责和边界，便于理解和维护。例如，Lexer负责词法分析，PrecedenceParser负责语法分析，SemanticAnalyzer负责语义分析，CodegenVisitor负责代码生成。',
        '（2）接口简洁：模块之间通过简洁的接口进行通信，降低了模块之间的耦合度。例如，Lexer输出Token序列，PrecedenceParser接收Token序列并输出AST，SemanticAnalyzer接收AST并输出四元式。',
        '（3）易于测试：每个模块可以独立进行测试，便于发现和定位问题。例如，可以单独测试词法分析器，验证它能否正确识别各种Token。',
        '（4）易于扩展：当需要添加新的功能时，只需修改或添加相应的模块，不会影响其他模块。例如，如果要支持新的语言特性，只需修改文法定义和语义分析器。',
    ]

    for advantage in modularity_advantages:
        add_paragraph(doc, advantage, indent=True)

    # 模块依赖关系表
    p = doc.add_paragraph()
    run = p.add_run('各模块之间的依赖关系如表4-2所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('表4-2 模块依赖关系')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['模块', '依赖模块', '被依赖模块', '主要接口']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    dependency_data = [
        ('Lexer', 'ErrorHandler, SymbolTable', 'PrecedenceParser', 'tokenize(), nextToken()'),
        ('Grammar', '无', 'PrecedenceParser', 'getRelation(), getProductions()'),
        ('PrecedenceParser', 'Grammar, ErrorHandler, SemanticAnalyzer', 'main', 'parse(), getAST()'),
        ('SemanticAnalyzer', '无', 'PrecedenceParser, CodegenVisitor', 'emit(), declareVariable()'),
        ('CodegenVisitor', 'SemanticAnalyzer', 'main', 'visit(), generate()'),
        ('AST', '无', 'PrecedenceParser, CodegenVisitor', 'ASTNode结构'),
        ('SymbolTable', '无', 'Lexer, SemanticAnalyzer', 'add(), lookup()'),
        ('ErrorHandler', '无', 'Lexer, PrecedenceParser', 'reportError(), hasErrors()'),
    ]

    for i, (col1, col2, col3, col4) in enumerate(dependency_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        table.rows[i].cells[3].text = col4
        for j in range(4):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 4.1.3 完整的编译流程
    add_heading(doc, '4.1.3 完整的编译流程', 3)

    p = doc.add_paragraph()
    run = p.add_run('本编译器实现了从源代码到目标代码的完整编译流程，包括词法分析、语法分析、语义分析和代码生成四个阶段。每个阶段都有明确的输入和输出，形成了完整的编译流水线。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('完整的编译流程具有以下特点：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    flow_features = [
        '（1）阶段性：编译过程被清晰地划分为四个阶段，每个阶段完成特定的任务，便于理解和实现。',
        '（2）流水线式：各阶段之间通过数据流连接，前一阶段的输出是后一阶段的输入，形成流水线式的处理过程。',
        '（3）可验证性：每个阶段的输出都可以被验证，便于发现和定位问题。例如，可以检查词法分析器输出的Token序列是否正确，语法分析器输出的AST是否正确等。',
        '（4）可扩展性：当需要支持新的语言特性时，只需修改相应的阶段，不会影响其他阶段。',
    ]

    for feature in flow_features:
        add_paragraph(doc, feature, indent=True)

    # 4.1.4 完善的错误处理
    add_heading(doc, '4.1.4 完善的错误处理', 3)

    p = doc.add_paragraph()
    run = p.add_run('本编译器实现了完善的错误处理机制，能够在编译过程中检测和报告各种错误。错误处理机制包括以下几个方面：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    error_handling = [
        '（1）词法错误处理：能够检测非法字符、数字格式错误、字符常量未闭合等词法错误，并给出错误位置和错误类型信息。',
        '（2）语法错误处理：能够检测语法错误，如缺少分号、括号不匹配、语句结构不正确等，并给出错误位置和错误类型信息。',
        '（3）语义错误处理：能够检测语义错误，如变量未声明、类型不匹配、函数参数错误等，并给出错误位置和错误类型信息。',
        '（4）错误恢复：在遇到错误时，编译器能够尝试恢复并继续编译后续代码，而不是立即终止，这样可以一次性报告多个错误。',
    ]

    for handling in error_handling:
        add_paragraph(doc, handling, indent=True)

    # 4.1.5 丰富的输出信息
    add_heading(doc, '4.1.5 丰富的输出信息', 3)

    p = doc.add_paragraph()
    run = p.add_run('本编译器在编译过程中会输出丰富的中间信息，便于调试和理解编译过程。输出信息包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    output_info = [
        '（1）Token序列：词法分析器输出的Token序列，保存在output/Lexer/tokens.txt文件中。',
        '（2）产生式列表：文法的所有产生式，保存在output/Parser/productions.txt文件中。',
        '（3）FIRSTVT和LASTVT集合：计算得到的FIRSTVT和LASTVT集合，保存在output/Parser/firstvt.txt和output/Parser/lastvt.txt文件中。',
        '（4）优先关系表：构建的优先关系表，保存在output/Parser/precedence_table.txt文件中。',
        '（5）语法分析过程：语法分析的详细过程，包括移进、归约、接受等操作，保存在output/Parser/trace.txt文件中。',
        '（6）语法树：构建的语法树和抽象语法树，保存在output/Parser/parse_tree.txt和output/Parser/ast.txt文件中。',
        '（7）四元式：生成的四元式中间代码，保存在output/Codegen/quads.txt文件中。',
        '（8）VM指令：生成的VM指令，保存在output/Codegen/Main.vm文件中。',
    ]

    for info in output_info:
        add_paragraph(doc, info, indent=True)

    # 输出文件汇总表
    p = doc.add_paragraph()
    run = p.add_run('编译器输出文件汇总如表4-3所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('表4-3 编译器输出文件汇总')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['输出文件', '内容说明', '生成阶段']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    output_files = [
        ('output/Lexer/tokens.txt', 'Token序列', '词法分析'),
        ('output/Parser/productions.txt', '产生式列表', '语法分析'),
        ('output/Parser/firstvt.txt', 'FIRSTVT集合', '语法分析'),
        ('output/Parser/lastvt.txt', 'LASTVT集合', '语法分析'),
        ('output/Parser/precedence_table.txt', '优先关系表', '语法分析'),
        ('output/Parser/trace.txt', '语法分析过程', '语法分析'),
        ('output/Parser/ast.txt', '抽象语法树', '语法分析'),
        ('output/Codegen/quads.txt', '四元式', '语义分析'),
    ]

    for i, (col1, col2, col3) in enumerate(output_files, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # ==================== 4.2 实训体会 ====================
    add_heading(doc, '4.2 实训体会', 2)

    p = doc.add_paragraph()
    run = p.add_run('通过本次实训，我们获得了宝贵的实践经验和深刻的体会，主要体现在以下几个方面：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 4.2.1 理论与实践的结合
    add_heading(doc, '4.2.1 理论与实践的结合', 3)

    p = doc.add_paragraph()
    run = p.add_run('本次实训最大的收获是将编译原理课程中学到的理论知识应用到了实际项目中。在课堂上，我们学习了词法分析、语法分析、语义分析、代码生成等编译器的基本原理，但这些知识往往停留在理论层面。通过本次实训，我们亲手实现了一个完整的编译器，将理论知识转化为了实际的代码和功能。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('在实现过程中，我们深刻体会到了理论与实践之间的差距。例如，简单优先法的理论看起来很简单，但在实际实现时，需要处理很多细节问题，如优先关系表的构建、移进-归约的具体实现、错误处理等。通过解决这些问题，我们对简单优先法有了更深入的理解。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('此外，我们还学到了很多课堂上没有涉及的知识，如抽象语法树的设计、符号表的管理、四元式的生成等。这些知识对于理解编译器的工作原理非常重要，也是本次实训的重要收获。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 4.2.2 工程实践能力的提升
    add_heading(doc, '4.2.2 工程实践能力的提升', 3)

    p = doc.add_paragraph()
    run = p.add_run('本次实训不仅是一个编译器开发项目，更是一次完整的软件工程实践。通过本次实训，我们的工程实践能力得到了显著提升，主要体现在以下几个方面：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    engineering_skills = [
        '（1）需求分析能力：在项目开始之前，我们需要对编译器的功能需求进行详细分析，明确编译器需要支持的语言特性和需要实现的功能。通过需求分析，我们学会了如何将模糊的需求转化为明确的功能规格。',
        '（2）系统设计能力：在需求分析的基础上，我们需要进行编译器的总体设计和详细设计，确定系统的架构、模块划分、接口定义等。通过系统设计，我们学会了如何将复杂系统分解为多个简单的模块，以及如何设计模块之间的接口。',
        '（3）编码实现能力：在系统设计的基础上，我们需要将设计转化为实际的代码。通过编码实现，我们学会了如何编写清晰、可维护的代码，如何使用数据结构和算法解决实际问题。',
        '（4）测试调试能力：在编码实现之后，我们需要对编译器进行测试和调试，发现和修复错误。通过测试调试，我们学会了如何设计测试用例，如何使用调试工具定位问题，如何分析和解决错误。',
        '（5）文档编写能力：在项目开发过程中，我们需要编写各种文档，如需求文档、设计文档、测试文档等。通过文档编写，我们学会了如何清晰地表达技术方案，如何记录和分享知识。',
    ]

    for skill in engineering_skills:
        add_paragraph(doc, skill, indent=True)

    # 4.2.3 问题解决能力的提升
    add_heading(doc, '4.2.3 问题解决能力的提升', 3)

    p = doc.add_paragraph()
    run = p.add_run('在编译器开发过程中，我们遇到了很多问题和挑战，通过解决这些问题，我们的问题解决能力得到了显著提升。以下是一些典型的问题和解决方法：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 问题与解决方案表
    p = doc.add_paragraph()
    run = p.add_run('表4-4 典型问题与解决方案')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['问题', '原因分析', '解决方案']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    problems = [
        ('优先关系冲突', '文法不满足简单优先文法条件', '修改文法，消除右部相同、ε产生式、相邻非终结符'),
        ('变量作用域错误', '符号表作用域管理不正确', '使用作用域栈，支持进入和退出作用域'),
        ('类型检查失败', '类型兼容性判断不准确', '完善类型转换规则，添加隐式类型转换'),
        ('代码生成错误', 'VM指令生成不正确', '仔细检查四元式到VM指令的映射规则'),
        ('测试结果不正确', '编译器存在Bug', '使用调试工具定位问题，逐步修复错误'),
    ]

    for i, (col1, col2, col3) in enumerate(problems, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('通过解决这些问题，我们学会了如何分析问题的根本原因，如何设计解决方案，如何验证解决方案的有效性。这些经验对于我们未来的学习和工作都非常有价值。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 4.2.4 团队协作的重要性
    add_heading(doc, '4.2.4 团队协作的重要性', 3)

    p = doc.add_paragraph()
    run = p.add_run('虽然本次实训是个人项目，但在开发过程中，我们深刻体会到了团队协作的重要性。在遇到问题时，我们通过查阅资料、请教老师和同学、在网上搜索解决方案等方式获取帮助。这些经历让我们认识到，在实际的软件开发中，团队协作是不可或缺的。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('此外，我们还学会了如何使用版本控制工具（Git）管理代码，如何编写清晰的注释和文档，如何设计易于理解的接口。这些技能对于团队协作非常重要，也是本次实训的重要收获。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 4.2.5 对计算机系统的深入理解
    add_heading(doc, '4.2.5 对计算机系统的深入理解', 3)

    p = doc.add_paragraph()
    run = p.add_run('通过本次实训，我们对计算机系统有了更深入的理解。编译器是计算机系统中最核心的系统软件之一，它连接了高级编程语言和底层硬件。通过实现编译器，我们深入理解了以下内容：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    understanding = [
        '（1）程序语言的本质：通过设计和实现mico C语言，我们深入理解了程序语言的语法、语义和语用，理解了语言设计的原则和方法。',
        '（2）编译器的工作原理：通过实现词法分析、语法分析、语义分析、代码生成等阶段，我们深入理解了编译器的工作原理和实现方法。',
        '（3）虚拟机技术：通过将编译器的目标代码在Hack虚拟机上运行，我们深入理解了虚拟机的架构和工作原理，理解了虚拟机与物理机之间的关系。',
        '（4）计算机系统的层次结构：通过本次实训，我们更加清晰地认识到计算机系统的层次结构，理解了各个层次之间的关系和作用。',
    ]

    for u in understanding:
        add_paragraph(doc, u, indent=True)

    # ==================== 4.3 改进建议 ====================
    add_heading(doc, '4.3 改进建议', 2)

    p = doc.add_paragraph()
    run = p.add_run('虽然本次实训成功实现了一个基本的编译器，但仍有一些可以改进的地方。以下是一些改进建议：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 4.3.1 功能扩展
    add_heading(doc, '4.3.1 功能扩展', 3)

    p = doc.add_paragraph()
    run = p.add_run('当前编译器支持的语言特性相对有限，可以从以下几个方面进行功能扩展：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    extensions = [
        '（1）支持更多数据类型：当前编译器只支持int、char、double三种基本类型，可以扩展支持字符串、布尔类型、枚举类型等。',
        '（2）支持更多运算符：当前编译器只支持基本的算术和比较运算符，可以扩展支持逻辑运算符（&&、||、!）、位运算符（&、|、^、~）、赋值运算符（+=、-=、*=、/=）等。',
        '（3）支持更多控制结构：当前编译器只支持if-else和while两种控制结构，可以扩展支持for循环、switch-case语句、do-while循环等。',
        '（4）支持函数特性：当前编译器的函数支持相对简单，可以扩展支持函数重载、默认参数、可变参数等高级特性。',
        '（5）支持面向对象特性：可以扩展支持类、对象、继承、多态等面向对象特性。',
    ]

    for ext in extensions:
        add_paragraph(doc, ext, indent=True)

    # 4.3.2 性能优化
    add_heading(doc, '4.3.2 性能优化', 3)

    p = doc.add_paragraph()
    run = p.add_run('当前编译器的性能还有提升空间，可以从以下几个方面进行性能优化：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    optimizations = [
        '（1）词法分析优化：可以使用更高效的词法分析算法，如DFA最小化、正则表达式优化等，提高词法分析的速度。',
        '（2）语法分析优化：可以使用更高效的语法分析算法，如LR分析法，提高语法分析的速度和能力。',
        '（3）代码优化：可以在代码生成阶段添加代码优化功能，如常量折叠、死代码消除、循环优化等，生成更高效的目标代码。',
        '（4）内存优化：可以优化内存管理，减少内存分配和释放的次数，提高编译器的内存使用效率。',
    ]

    for opt in optimizations:
        add_paragraph(doc, opt, indent=True)

    # 4.3.3 错误处理改进
    add_heading(doc, '4.3.3 错误处理改进', 3)

    p = doc.add_paragraph()
    run = p.add_run('当前编译器的错误处理机制还有改进空间，可以从以下几个方面进行改进：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    error_improvements = [
        '（1）更友好的错误提示：当前编译器的错误提示相对简单，可以改进为更友好的错误提示，包括错误位置、错误类型、错误原因、修复建议等信息。',
        '（2）更好的错误恢复：当前编译器的错误恢复机制相对简单，可以改进为更智能的错误恢复机制，能够在遇到错误时自动修复或跳过错误，继续编译后续代码。',
        '（3）更全面的错误检测：当前编译器的错误检测能力还有限，可以扩展支持更多的错误类型检测，如未使用的变量、不可达代码、潜在的类型错误等。',
        '（4）错误信息国际化：可以支持多语言的错误信息，方便不同语言的用户使用。',
    ]

    for imp in error_improvements:
        add_paragraph(doc, imp, indent=True)

    # 4.3.4 用户体验改进
    add_heading(doc, '4.3.4 用户体验改进', 3)

    p = doc.add_paragraph()
    run = p.add_run('当前编译器的用户体验还有改进空间，可以从以下几个方面进行改进：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    ux_improvements = [
        '（1）图形用户界面：当前编译器是命令行工具，可以开发图形用户界面，提供更友好的用户交互体验。',
        '（2）集成开发环境：可以开发集成开发环境（IDE），提供代码编辑、语法高亮、自动补全、错误提示等功能。',
        '（3）调试工具：可以开发调试工具，支持断点设置、单步执行、变量查看等功能，方便用户调试程序。',
        '（4）文档生成：可以自动生成API文档、用户手册等文档，方便用户使用编译器。',
    ]

    for imp in ux_improvements:
        add_paragraph(doc, imp, indent=True)

    # 4.3.5 代码质量改进
    add_heading(doc, '4.3.5 代码质量改进', 3)

    p = doc.add_paragraph()
    run = p.add_run('当前编译器的代码质量还有改进空间，可以从以下几个方面进行改进：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    code_improvements = [
        '（1）代码重构：可以对现有代码进行重构，提高代码的可读性和可维护性。例如，可以提取公共函数，消除代码重复，简化复杂的逻辑。',
        '（2）单元测试：可以添加更全面的单元测试，提高代码的可靠性和稳定性。',
        '（3）代码审查：可以进行代码审查，发现和修复潜在的问题，提高代码质量。',
        '（4）文档完善：可以完善代码注释和文档，提高代码的可理解性。',
    ]

    for imp in code_improvements:
        add_paragraph(doc, imp, indent=True)

    # 改进建议汇总表
    p = doc.add_paragraph()
    run = p.add_run('改进建议汇总如表4-5所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('表4-5 改进建议汇总')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['改进方向', '改进内容', '优先级']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    improvements_summary = [
        ('功能扩展', '支持更多数据类型、运算符、控制结构', '高'),
        ('性能优化', '词法分析、语法分析、代码优化', '中'),
        ('错误处理', '更友好的错误提示、更好的错误恢复', '高'),
        ('用户体验', '图形界面、IDE、调试工具', '低'),
        ('代码质量', '代码重构、单元测试、文档完善', '中'),
    ]

    for i, (col1, col2, col3) in enumerate(improvements_summary, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # ==================== 4.4 总结 ====================
    add_heading(doc, '4.4 总结', 2)

    p = doc.add_paragraph()
    run = p.add_run('本次实训历时四周，我们成功设计并实现了一个基于简单优先法的微型程序语言编译器。该编译器能够将mico C语言源程序编译成Hack虚拟机的VM指令，并在虚拟机上正确运行。通过本次实训，我们取得了以下主要成果：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    achievements = [
        '（1）设计并实现了mico C语言的词法分析器，能够正确识别关键字、标识符、常量、运算符、分隔符等各种Token。',
        '（2）设计并实现了基于简单优先法的语法分析器，能够正确进行语法分析，构建语法树和抽象语法树。',
        '（3）设计并实现了语义分析器，能够进行类型检查、符号表管理、中间代码生成等语义处理。',
        '（4）设计并实现了代码生成器，能够将四元式中间代码转换成Hack虚拟机的VM指令。',
        '（5）通过两个测试用例（简单算术运算和冒泡排序）的验证，证明编译器的各个模块都能正确工作，编译器的整体功能满足设计要求。',
    ]

    for achievement in achievements:
        add_paragraph(doc, achievement, indent=True)

    p = doc.add_paragraph()
    run = p.add_run('通过本次实训，我们不仅掌握了编译器设计与实现的基本方法，还提升了工程实践能力、问题解决能力和团队协作能力。这些经验和能力对于我们未来的学习和工作都非常有价值。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('最后，感谢指导教师的悉心指导和帮助，感谢同学们的支持和鼓励。本次实训是一次非常有意义的学习经历，我们将继续努力，不断提升自己的技术水平和工程能力。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 项目成果汇总表
    p = doc.add_paragraph()
    run = p.add_run('项目成果汇总如表4-6所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('表4-6 项目成果汇总')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['项目', '成果']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    results = [
        ('源语言', 'mico C（微型C语言）'),
        ('目标语言', 'Hack虚拟机VM指令'),
        ('语法分析方法', '简单优先法'),
        ('中间代码', '四元式'),
        ('测试用例', '2个（简单算术运算、冒泡排序）'),
        ('测试结果', '全部通过'),
        ('代码行数', '约3000行C++代码'),
    ]

    for i, (col1, col2) in enumerate(results, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        for j in range(2):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 保存文档
    doc.save('D:\\c_cpp\\ShiXvn_Compiler\\report\\第四章_实训小结.docx')
    print('第四章生成成功！')

if __name__ == '__main__':
    create_chapter4()
