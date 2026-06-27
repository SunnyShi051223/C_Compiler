# -*- coding: utf-8 -*-
"""
生成课程报告Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def create_report():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置行间距
    style.paragraph_format.line_spacing = 1.5

    # ==================== 封面 ====================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('计算机系统能力实训报告')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('基于简单优先法的微型程序语言编译器设计与实现')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for _ in range(8):
        doc.add_paragraph()

    # 信息表格
    info_items = [
        ('学生姓名：', 'XXXX'),
        ('专业班级：', 'XXXXXX'),
        ('指导教师：', 'XXXX'),
        ('工作单位：', '计算机与人工智能学院'),
        ('完成日期：', '2026年6月'),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(value)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.underline = True

    # 分页
    doc.add_page_break()

    # ==================== 目录页 ====================
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('目  录')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    toc_items = [
        '第一章  实训任务概述',
        '  1.1  任务背景',
        '  1.2  任务要求',
        '  1.3  任务内容',
        '  1.4  开发环境',
        '  1.5  本章小结',
        '第二章  系统分析',
        '  2.1  源语言定义',
        '  2.2  目标语言定义',
        '  2.3  编译器功能需求',
        '第三章  系统开发',
        '  3.1  总体设计',
        '  3.2  详细设计',
        '  3.3  编码实现',
        '  3.4  测试与验证',
        '第四章  实训小结',
        '  4.1  设计特点',
        '  4.2  实训体会',
        '  4.3  改进建议',
        '  4.4  总结',
        '参考文献',
    ]

    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ==================== 第一章 ====================
    # 章标题
    h1 = doc.add_heading('第一章  实训任务概述', level=1)
    for run in h1.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 1.1 任务背景
    h2 = doc.add_heading('1.1 任务背景', level=2)
    for run in h2.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    h3 = doc.add_heading('1.1.1 实训目的', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('计算机系统能力实训是计算机科学与技术专业教学计划中的重要实践教学环节，其目的是进一步加强学生对计算机组成原理、操作系统、编译原理、汇编语言等专业核心课程知识的理解和掌握。通过本次为期四周的综合性实训，学生需要将前期所学的理论知识进行综合运用，设计并实现一个计算机系统软件——高级编程语言编译器，从而培养计算机系统理论和技术的应用能力，提高学生的工程实践能力和创新意识。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('编译器是计算机系统中最核心的系统软件之一，它承担着将高级编程语言编写的源程序翻译成计算机能够直接执行的机器代码的重要任务。编译器的设计与实现涉及形式语言与自动机理论、数据结构与算法、程序设计语言原理等多个计算机科学领域的知识，是计算机科学理论与实践相结合的典型应用。通过编译器的设计与实现，学生可以深入理解计算机系统的工作原理，掌握软件系统设计的基本方法，提高分析问题和解决问题的能力。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h3 = doc.add_heading('1.1.2 选题意义', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('本次实训选择"基于简单优先法的微型程序语言编译器设计与实现"作为题目，具有以下几方面的意义：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    meanings = [
        '第一，简单优先法是一种经典的自底向上语法分析方法，它通过预先计算文法符号之间的优先关系来指导语法分析过程。与LL分析法和LR分析法相比，简单优先法的原理更加直观，实现相对简单，适合作为编译原理课程的实践教学内容。通过实现基于简单优先法的编译器，学生可以深入理解自底向上语法分析的基本原理和实现方法。',
        '第二，微型程序语言（mico C）是一种简化的C语言子集，它保留了C语言的核心语法特性，包括变量声明、赋值语句、控制结构、函数定义等，同时简化了部分复杂的语法特性。通过设计和实现mico C语言的编译器，学生可以掌握程序语言设计的基本方法，理解编译器各个阶段的工作原理。',
        '第三，本项目选择Hack虚拟机作为目标平台，使用VM指令集作为目标语言。Hack虚拟机是nand2tetris课程中的重要组成部分，它提供了一个简洁而完整的虚拟机架构。通过将编译器的目标代码在Hack虚拟机上运行，学生可以验证编译器的正确性，同时加深对虚拟机技术的理解。'
    ]

    for meaning in meanings:
        p = doc.add_paragraph()
        run = p.add_run(meaning)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h3 = doc.add_heading('1.1.3 实训要求', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('根据实训任务书的要求，编译器软件需按照软件工程的思想，完成系统软件的分析、设计及实现，翻译的目标代码能够在虚拟目标机上正确运行。具体而言，本次实训需要完成以下主要任务：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    requirements = [
        '（1）源语言和目标语言的定义：采用符合编译器要求的文法定义一个结构完整的源语言程序，包含输入输出语句、赋值语句、循环结构的语句、选择结构的语句等，确定虚拟机及目标代码指令集。',
        '（2）词法分析程序的设计及实现：根据源语言的词法规则，设计及实现源语言的词法分析程序，完成源程序单词的截取，识别词法错误，并转换成特征字序列，为语法分析做准备。',
        '（3）语法分析程序设计及实现：选择简单优先语法分析方法，理解其主要思想，设计相关的分析表和语法分析步骤，完成源程序是否合法的判断，并指出其中的语法错误，具备容错功能。',
        '（4）目标代码生成：设计语义翻译的属性文法，在语法分析的基础上，采用语法制导的方法实现目标代码的生成。',
        '（5）目标代码的运行：利用已有虚拟机的解析器运行目标程序，得到与源程序完全一致的结果。通过多个测试用例验证编译器的正确性。'
    ]

    for req in requirements:
        p = doc.add_paragraph()
        run = p.add_run(req)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 1.2 任务要求
    h2 = doc.add_heading('1.2 任务要求', level=2)
    for run in h2.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    h3 = doc.add_heading('1.2.1 软件工程要求', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('本次实训要求按照软件工程的思想进行编译器的开发，具体包括以下几个方面：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    sw_requirements = [
        '需求分析阶段：在项目开始之前，需要对编译器的功能需求和非功能需求进行详细分析。功能需求包括编译器需要支持的语言特性、输入输出格式、错误处理机制等；非功能需求包括性能要求、可维护性要求、可扩展性要求等。通过需求分析，明确编译器的边界和目标，为后续的设计和开发工作奠定基础。',
        '系统设计阶段：在需求分析的基础上，进行编译器的总体设计和详细设计。总体设计确定编译器的整体架构、模块划分、接口定义等；详细设计对每个模块的算法、数据结构、实现细节进行详细说明。系统设计阶段需要输出设计文档，作为编码实现的依据。',
        '编码实现阶段：按照设计文档进行编码实现，要求代码结构清晰、注释完善、符合编码规范。在编码过程中，需要注意模块之间的解耦，提高代码的可重用性和可维护性。',
        '测试验证阶段：通过多个测试用例验证编译器的正确性。测试用例需要覆盖编译器的各个功能模块，包括正常情况和异常情况。测试过程中需要记录测试结果，分析和解决发现的问题。'
    ]

    for req in sw_requirements:
        p = doc.add_paragraph()
        # 添加加粗的阶段名称
        colon_idx = req.index('：')
        run = p.add_run(req[:colon_idx+1])
        run.font.bold = True
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(req[colon_idx+1:])
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h3 = doc.add_heading('1.2.2 功能要求', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('编译器需要支持以下语言特性，如表1-1所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表1-1 编译器支持的语言特性
    p = doc.add_paragraph()
    run = p.add_run('表1-1 编译器支持的语言特性')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=19, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ['类别', '语言特性', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    # 表格数据
    data = [
        ('数据类型', 'int', '整数类型'),
        ('', 'char', '字符类型'),
        ('', 'double', '双精度浮点类型'),
        ('', '数组', '一维数组声明和访问'),
        ('', '结构体', '结构体定义和成员访问'),
        ('变量操作', '变量声明', '声明变量并可选初始化'),
        ('', '赋值语句', '变量赋值操作'),
        ('算术运算', '加法（+）', '两个操作数相加'),
        ('', '减法（-）', '两个操作数相减'),
        ('', '乘法（*）', '两个操作数相乘'),
        ('', '除法（/）', '两个操作数相除'),
        ('', '取模（%）', '求余数'),
        ('控制结构', 'if-else', '条件分支语句'),
        ('', 'while', '循环语句'),
        ('输入输出', 'read', '读取输入'),
        ('', 'write', '输出结果'),
        ('函数', '函数定义', '定义函数'),
        ('', '函数调用', '调用函数'),
    ]

    for i, (col1, col2, col3) in enumerate(data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 1.3 任务内容
    h2 = doc.add_heading('1.3 任务内容', level=2)
    for run in h2.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    h3 = doc.add_heading('1.3.1 源语言和目标语言的定义', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('源语言定义')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('本项目定义的源语言称为mico C（微型C语言），它是一种简化的C语言子集。mico C语言保留了C语言的核心语法特性，包括数据类型（int、char、double、数组、结构体）、变量声明与赋值、算术运算、比较运算、逻辑运算、控制结构（if-else、while）、输入输出（read、write）、函数定义与调用等。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('目标语言定义')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('本项目选择Hack虚拟机作为目标平台，使用VM指令集作为目标语言。Hack虚拟机是一种基于栈的虚拟机，其指令集包括栈操作指令（push、pop）、算术逻辑指令（add、sub、neg、eq、gt、lt、and、or、not）、流程控制指令（label、goto、if-goto）、函数调用指令（function、call、return）等。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h3 = doc.add_heading('1.3.2 词法分析程序的设计及实现', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('词法分析是编译器的第一个阶段，其主要任务是将源程序的字符序列转换成Token序列。Token是词法分析的基本单位，每个Token包含两个信息：Token类型和Token值。本次实训中，词法分析程序需要完成单词截取、词法错误识别、Token序列生成等工作，为语法分析做准备。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h3 = doc.add_heading('1.3.3 语法分析程序设计及实现', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('语法分析是编译器的第二个阶段，其主要任务是根据语言的语法规则，分析Token序列的语法结构，判断源程序是否符合语法规则。本次实训选择简单优先语法分析方法，这是一种经典的自底向上语法分析方法。简单优先法通过预先计算文法符号之间的优先关系来指导语法分析过程，具有原理直观、实现简单的特点。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h3 = doc.add_heading('1.3.4 目标代码生成', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('目标代码生成是编译器的最后一个阶段，其主要任务是将语法分析和语义分析的结果转换成目标代码。本次实训中，目标代码生成采用语法制导翻译的方法，在文法的产生式中嵌入语义动作，当产生式被归约时执行相应的语义动作，生成符合Hack虚拟机规范的VM指令。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h3 = doc.add_heading('1.3.5 目标代码的运行', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('目标代码的运行是验证编译器正确性的重要环节。编译器生成的VM代码经过VMTranslator转换为汇编代码，再经过Assembler转换为机器码，最终在CPUEmulator或VMEmulator上执行。通过多个测试用例验证编译器的正确性，确保目标代码的执行结果与源程序的语义完全一致。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 1.4 开发环境
    h2 = doc.add_heading('1.4 开发环境', level=2)
    for run in h2.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    h3 = doc.add_heading('1.4.1 硬件环境', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('本次实训使用的硬件环境如表1-4所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表1-4 硬件环境配置
    p = doc.add_paragraph()
    run = p.add_run('表1-4 硬件环境配置')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['项目', '配置', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    hw_data = [
        ('处理器', 'Intel Core i5-12400', '6核12线程，主频2.5GHz'),
        ('内存', '16GB DDR4', '满足编译器运行需求'),
        ('硬盘', '512GB SSD', '提供足够的存储空间'),
        ('显示器', '1920×1080分辨率', '提供良好的开发环境'),
    ]

    for i, (col1, col2, col3) in enumerate(hw_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    h3 = doc.add_heading('1.4.2 软件环境', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('本次实训使用的软件环境如表1-5所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表1-5 软件环境配置
    p = doc.add_paragraph()
    run = p.add_run('表1-5 软件环境配置')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['类别', '软件', '版本', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    sw_data = [
        ('操作系统', 'Windows', '11', '开发和运行环境'),
        ('编程语言', 'C++', '17', '编译器实现语言'),
        ('编译器', 'GCC/G++', '13.2.0', 'C++编译器'),
        ('构建工具', 'CMake', '3.28', '跨平台构建工具'),
        ('开发工具', 'Visual Studio Code', '1.85', '代码编辑器'),
        ('版本控制', 'Git', '2.43', '版本控制系统'),
        ('调试工具', 'GDB', '13.2', '程序调试工具'),
    ]

    for i, (col1, col2, col3, col4) in enumerate(sw_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        table.rows[i].cells[3].text = col4
        for j in range(4):
            for paragraph in table.rows[i].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    h3 = doc.add_heading('1.4.3 目标平台', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('本次实训的目标平台如表1-6所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表1-6 目标平台配置
    p = doc.add_paragraph()
    run = p.add_run('表1-6 目标平台配置')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['项目', '说明', '用途']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    target_data = [
        ('目标虚拟机', 'Hack虚拟机', '运行目标代码'),
        ('指令集', 'VM指令集', '目标代码格式'),
        ('测试工具', 'CPUEmulator', '硬件级模拟器'),
        ('测试工具', 'VMEmulator', '虚拟机级模拟器'),
        ('汇编工具', 'Assembler', '汇编代码转换'),
    ]

    for i, (col1, col2, col3) in enumerate(target_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    h3 = doc.add_heading('1.4.4 工具链', level=3)
    for run in h3.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('编译器工具链由多个工具组成，各工具之间的关系如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('源程序(.txt) → ShiXvn_Compiler → VM代码(.vm) → VMTranslator → 汇编代码(.asm) → Assembler → 机器码(.hack) → CPUEmulator/VMEmulator → 执行结果')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('【图表说明】请绘制一个工具链流程图，展示从源代码到最终执行的完整流程。图中需要包含源程序文件（.txt）、编译器（ShiXvn_Compiler）、VM代码文件（.vm）、VMTranslator工具、汇编代码文件（.asm）、Assembler工具、机器码文件（.hack）、CPUEmulator或VMEmulator、执行结果，各个阶段之间用箭头连接，标注转换关系。')
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 1.5 本章小结
    h2 = doc.add_heading('1.5 本章小结', level=2)
    for run in h2.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    run = p.add_run('本章对基于简单优先法的微型程序语言编译器设计与实现的实训任务进行了全面的概述。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('首先，在任务背景部分，介绍了本次实训的目的和意义，说明了编译器在计算机系统中的重要地位，以及选择简单优先法和mico C语言的原因。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('其次，在任务要求部分，详细说明了软件工程要求、功能要求、目标代码要求和文档要求，明确了本次实训需要完成的具体任务。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('然后，在任务内容部分，详细介绍了源语言和目标语言的定义、词法分析程序的设计及实现、语法分析程序设计及实现、目标代码生成、目标代码的运行等各个阶段的工作内容。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('最后，在开发环境部分，介绍了本次实训使用的硬件环境、软件环境、目标平台、工具链和项目目录结构，为后续章节的详细设计和实现奠定了基础。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('通过本章的介绍，读者可以对本次实训任务有一个全面的了解，明确编译器需要实现的功能和达到的目标，为后续章节的学习提供指导。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存文档
    doc.save('D:\\c_cpp\\ShiXvn_Compiler\\report\\课程报告_基于简单优先法的微型程序语言编译器设计与实现.docx')
    print('报告生成成功！')

if __name__ == '__main__':
    create_report()
