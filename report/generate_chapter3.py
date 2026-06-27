# -*- coding: utf-8 -*-
"""
生成第三章 系统开发 Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def add_heading(doc, text, level):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_paragraph(doc, text, bold_prefix=None, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(text)
    else:
        run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.left_indent = Pt(36)
    return p

def add_code(doc, code_lines):
    """添加代码块"""
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10.5)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(72)
        # 设置段落背景色
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    return p

def add_note(doc, text):
    """添加图表说明（红色斜体）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def create_chapter3():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ==================== 第三章标题 ====================
    add_heading(doc, '第三章  系统开发', 1)

    p = doc.add_paragraph()
    run = p.add_run('系统开发是编译器实现的核心阶段，主要包括总体设计、详细设计、编码实现和测试验证四个部分。本章将详细介绍编译器的架构设计、各模块的算法设计与实现、核心代码实现，以及测试用例和测试结果。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 3.1 总体设计 ====================
    add_heading(doc, '3.1 总体设计', 2)

    # 3.1.1 编译器架构设计
    add_heading(doc, '3.1.1 编译器架构设计', 3)

    p = doc.add_paragraph()
    run = p.add_run('本编译器采用经典的多阶段编译架构，将编译过程分为词法分析、语法分析、语义分析和代码生成四个阶段。每个阶段都有明确的输入和输出，模块之间通过接口进行通信。编译器的整体架构如图3-1所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_note(doc, '【图表说明】请绘制一个编译器整体架构图，展示以下内容：')
    add_note(doc, '1. 顶层：源程序输入 → 编译器 → 目标代码输出')
    add_note(doc, '2. 编译器内部模块：词法分析器（Lexer）、语法分析器（PrecedenceParser）、语义分析器（SemanticAnalyzer）、代码生成器（CodegenVisitor）')
    add_note(doc, '3. 辅助模块：符号表（SymbolTable）、错误处理器（ErrorHandler）、文法处理器（Grammar）')
    add_note(doc, '4. 数据流：源程序 → Token序列 → 语法树/AST → 四元式 → VM指令')
    add_note(doc, '5. 使用不同颜色区分各个模块，用箭头表示数据流向')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('编译器架构采用模块化设计，各模块职责清晰，接口简洁。主要模块包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 模块说明表
    p = doc.add_paragraph()
    run = p.add_run('表3-1 编译器模块说明')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['模块名称', '类名', '主要职责']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    modules = [
        ('词法分析器', 'Lexer', '将源程序字符序列转换为Token序列'),
        ('语法分析器', 'PrecedenceParser', '使用简单优先法进行语法分析，构建语法树'),
        ('语义分析器', 'SemanticAnalyzer', '进行语义检查，生成四元式中间代码'),
        ('代码生成器', 'CodegenVisitor', '遍历AST，生成VM指令'),
        ('文法处理器', 'Grammar', '定义文法，计算FIRSTVT/LASTVT，构建优先关系表'),
        ('符号表', 'SymbolTable', '管理标识符信息'),
        ('错误处理器', 'ErrorHandler', '收集和报告编译错误'),
        ('抽象语法树', 'AST', '表示程序的语法结构'),
    ]

    for i, (col1, col2, col3) in enumerate(modules, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 2 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 3.1.2 编译流程设计
    add_heading(doc, '3.1.2 编译流程设计', 3)

    p = doc.add_paragraph()
    run = p.add_run('编译器的编译流程如图3-2所示，各阶段的功能和数据流如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_note(doc, '【图表说明】请绘制一个详细的编译流程图，展示以下内容：')
    add_note(doc, '1. 阶段一：词法分析 - 输入：源程序字符序列，输出：Token序列，工具：Lexer')
    add_note(doc, '2. 阶段二：语法分析 - 输入：Token序列，输出：语法树/AST，工具：PrecedenceParser')
    add_note(doc, '3. 阶段三：语义分析 - 输入：AST，输出：四元式，工具：SemanticAnalyzer')
    add_note(doc, '4. 阶段四：代码生成 - 输入：四元式，输出：VM指令，工具：CodegenVisitor')
    add_note(doc, '5. 各阶段之间用箭头连接，标注输入输出数据类型')

    doc.add_paragraph()

    # 编译流程表
    p = doc.add_paragraph()
    run = p.add_run('表3-2 编译流程说明')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['阶段', '模块', '输入', '输出', '主要工作']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    flow_data = [
        ('阶段一', 'Lexer', '源程序字符', 'Token序列', '单词截取、词法检查'),
        ('阶段二', 'PrecedenceParser', 'Token序列', '语法树/AST', '语法分析、构建语法树'),
        ('阶段三', 'SemanticAnalyzer', 'AST', '四元式', '语义检查、生成中间代码'),
        ('阶段四', 'CodegenVisitor', '四元式', 'VM指令', '指令选择、代码生成'),
    ]

    for i, (col1, col2, col3, col4, col5) in enumerate(flow_data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        table.rows[i].cells[3].text = col4
        table.rows[i].cells[4].text = col5
        for j in range(5):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 3.1.3 数据结构设计
    add_heading(doc, '3.1.3 数据结构设计', 3)

    p = doc.add_paragraph()
    run = p.add_run('编译器使用的主要数据结构包括Token、AST节点、符号表项、四元式等。下面分别介绍这些数据结构的设计。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Token数据结构
    add_paragraph(doc, 'Token数据结构', bold_prefix='（1）')

    p = doc.add_paragraph()
    run = p.add_run('Token是词法分析的基本单位，包含类型和值两个主要属性。Token的数据结构定义如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'struct Token {',
        '    TokenType type;      // Token类型',
        '    std::string value;   // Token值',
        '    int line;            // 行号',
        '    int column;          // 列号',
        '};',
    ])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('TokenType枚举定义了Token的类型，包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'enum class TokenType {',
        '    KEYWORD,        // 关键字: int, char, if, while...',
        '    IDENTIFIER,     // 标识符: 变量名、函数名',
        '    NUMBER,         // 整型常数: 42, 1024',
        '    CHAR_LITERAL,   // 字符常数: \'A\', \'\\n\'',
        '    OPERATOR,       // 运算符: +, -, *, =, ==',
        '    DELIMITER,      // 界符: (, ), {, }, ;',
        '    UNKNOWN,        // 未知/非法字符',
        '    END_OF_FILE     // 文件结束',
        '};',
    ])

    doc.add_paragraph()

    # AST节点数据结构
    add_paragraph(doc, 'AST节点数据结构', bold_prefix='（2）')

    p = doc.add_paragraph()
    run = p.add_run('抽象语法树（AST）是语法分析的输出，表示程序的语法结构。AST节点的数据结构定义如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'struct ASTNode {',
        '    ASTNodeType type;                              // 节点类型',
        '    std::string value;                             // 节点值',
        '    std::vector<std::shared_ptr<ASTNode>> children; // 子节点列表',
        '    int line;                                      // 行号',
        '};',
    ])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('ASTNodeType枚举定义了AST节点的类型，包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # AST节点类型表
    p = doc.add_paragraph()
    run = p.add_run('表3-3 AST节点类型说明')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=15, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['节点类型', '枚举值', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    ast_types = [
        ('程序', 'Program', '程序根节点，子节点为语句列表'),
        ('变量声明', 'VarDecl', '变量声明，value=类型'),
        ('结构体定义', 'StructDecl', '结构体定义，value=结构体名'),
        ('赋值语句', 'AssignStmt', '赋值语句，children=[LValue, Expr]'),
        ('if语句', 'IfStmt', '条件语句，children=[Cond, Then, Else]'),
        ('while语句', 'WhileStmt', '循环语句，children=[Cond, Body]'),
        ('return语句', 'ReturnStmt', '返回语句，children=[Expr]'),
        ('语句块', 'Block', '语句块，children=[Stmt, ...]'),
        ('二元表达式', 'BinaryExpr', '二元运算，value=运算符'),
        ('函数调用', 'CallExpr', '函数调用，value=函数名'),
        ('标识符', 'IdLeaf', '标识符叶子节点，value=变量名'),
        ('数字常量', 'NumLeaf', '数字常量叶子节点，value=数字'),
        ('字符常量', 'CharLeaf', '字符常量叶子节点，value=字符'),
        ('左值', 'LValueId', '左值标识符，value=变量名'),
    ]

    for i, (col1, col2, col3) in enumerate(ast_types, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        for j in range(3):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 2 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 符号表数据结构
    add_paragraph(doc, '符号表数据结构', bold_prefix='（3）')

    p = doc.add_paragraph()
    run = p.add_run('符号表用于存储标识符的信息，包括变量名、类型、偏移量等。符号表项的数据结构定义如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'struct SymbolEntry {',
        '    std::string name;    // 变量名',
        '    std::string type;    // 类型',
        '    int offset;          // 偏移量',
        '    bool isInitialized;  // 是否已初始化',
        '};',
    ])

    doc.add_paragraph()

    # 四元式数据结构
    add_paragraph(doc, '四元式数据结构', bold_prefix='（4）')

    p = doc.add_paragraph()
    run = p.add_run('四元式是中间代码的一种形式，包含四个部分：运算符、操作数1、操作数2、结果。四元式的数据结构定义如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'struct Quadruple {',
        '    std::string op;    // 运算符',
        '    std::string arg1;  // 操作数1',
        '    std::string arg2;  // 操作数2',
        '    std::string result;// 结果',
        '};',
    ])

    doc.add_paragraph()

    # 四元式示例表
    p = doc.add_paragraph()
    run = p.add_run('表3-4 四元式示例')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['编号', 'op', 'arg1', 'arg2', 'result']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    quad_examples = [
        ('1', '=', '10', '-', 'a'),
        ('2', '=', '20', '-', 'b'),
        ('3', '+', 'a', 'b', 't1'),
        ('4', '=', 't1', '-', 'c'),
        ('5', 'call', 'write', '1', 't2'),
        ('6', '=', 't2', '-', '-'),
    ]

    for i, (col1, col2, col3, col4, col5) in enumerate(quad_examples, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        table.rows[i].cells[3].text = col4
        table.rows[i].cells[4].text = col5
        for j in range(5):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10.5)

    doc.add_paragraph()

    # ==================== 3.2 详细设计 ====================
    add_heading(doc, '3.2 详细设计', 2)

    # 3.2.1 词法分析器设计
    add_heading(doc, '3.2.1 词法分析器设计', 3)

    add_paragraph(doc, '设计原理', bold_prefix='1. ')

    p = doc.add_paragraph()
    run = p.add_run('词法分析器采用有限自动机（DFA）的理论进行设计。有限自动机是一种识别模式的数学模型，它由状态集合、输入字母表、转移函数、初始状态和接受状态组成。词法分析器通过状态转换图来描述单词的识别过程。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_note(doc, '【图表说明】请绘制一个词法分析器的状态转换图，展示以下状态和转移：')
    add_note(doc, '1. 初始状态 → 根据输入字符转移到不同状态')
    add_note(doc, '2. 标识符/关键字状态：读取字母或数字')
    add_note(doc, '3. 数字状态：读取数字')
    add_note(doc, '4. 运算符状态：读取运算符')
    add_note(doc, '5. 字符常量状态：读取单引号内的字符')
    add_note(doc, '6. 使用圆圈表示状态，箭头表示转移，标注输入字符')

    doc.add_paragraph()

    add_paragraph(doc, '功能实现', bold_prefix='2. ')

    p = doc.add_paragraph()
    run = p.add_run('词法分析器的主要功能包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    lexer_functions = [
        '（1）跳过空白字符和注释：在识别单词之前，先跳过空格、制表符、换行符等空白字符。',
        '（2）识别关键字：将读取的标识符与关键字集合进行匹配，判断是否为关键字。',
        '（3）读取数字：读取连续的数字字符，生成NUMBER类型的Token。',
        '（4）读取字符常量：读取单引号括起来的字符，生成CHAR_LITERAL类型的Token。',
        '（5）读取标识符：读取以字母或下划线开头的字母数字序列，生成IDENTIFIER类型的Token。',
        '（6）读取运算符：读取运算符字符，生成OPERATOR类型的Token。',
        '（7）读取分隔符：读取分隔符字符，生成DELIMITER类型的Token。',
    ]

    for func in lexer_functions:
        add_paragraph(doc, func, indent=True)

    add_paragraph(doc, '关键算法', bold_prefix='3. ')

    p = doc.add_paragraph()
    run = p.add_run('词法分析器的核心算法是nextToken()函数，它从源程序中读取下一个Token。算法流程如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'Token Lexer::nextToken() {',
        '    // 1. 跳过空白字符和注释',
        '    skipWhitespaceAndComments();',
        '',
        '    // 2. 检查是否到达文件末尾',
        '    if (isAtEnd()) return {END_OF_FILE, "", line_, column_};',
        '',
        '    // 3. 根据当前字符类型调用相应的处理函数',
        '    char c = currentChar();',
        '    if (isDigit(c)) return readNumber();      // 读取数字',
        '    if (isAlpha(c)) return readIdentifierOrKeyword(); // 读取标识符/关键字',
        '    if (c == \'\\\'\') return readCharLiteral();   // 读取字符常量',
        '    if (isOperator(c)) return readOperator();  // 读取运算符',
        '    if (isDelimiter(c)) return readDelimiter();// 读取分隔符',
        '',
        '    // 4. 未知字符，报告错误',
        '    errorHandler_.reportError(...);',
        '    return {UNKNOWN, string(1, c), line_, column_};',
        '}',
    ])

    doc.add_paragraph()

    # 3.2.2 语法分析器设计
    add_heading(doc, '3.2.2 语法分析器设计（简单优先法）', 3)

    add_paragraph(doc, '设计原理', bold_prefix='1. ')

    p = doc.add_paragraph()
    run = p.add_run('简单优先法是一种经典的自底向上语法分析方法，其基本思想是：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    principles = [
        '（1）为文法中的所有符号对预先计算优先关系（小于<·、等于=、大于·>三种关系）。',
        '（2）在语法分析过程中，根据栈顶符号和当前输入符号的优先关系，决定是进行移进操作还是归约操作。',
        '（3）如果栈顶符号的优先级小于当前输入符号（<·），则进行移进操作，将当前输入符号压入栈中。',
        '（4）如果栈顶符号的优先级大于当前输入符号（·>），则进行归约操作，将栈顶的某个子串归约为对应的非终结符。',
        '（5）如果栈顶符号的优先级等于当前输入符号（=），则进行移进操作。',
    ]

    for principle in principles:
        add_paragraph(doc, principle, indent=True)

    add_paragraph(doc, 'FIRSTVT和LASTVT集合计算', bold_prefix='2. ')

    p = doc.add_paragraph()
    run = p.add_run('FIRSTVT(A)集合定义为：A能推导出的所有句型中，第一个终结符的集合。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('LASTVT(A)集合定义为：A能推导出的所有句型中，最后一个终结符的集合。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('FIRSTVT集合的计算规则：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        '规则1: 若有产生式 A → a... 或 A → Ba...',
        '       则 a ∈ FIRSTVT(A)',
        '',
        '规则2: 若有产生式 A → B...',
        '       且 a ∈ FIRSTVT(B)',
        '       则 a ∈ FIRSTVT(A)',
    ])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('LASTVT集合的计算规则：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        '规则1: 若有产生式 A → ...a 或 A → ...aB',
        '       则 a ∈ LASTVT(A)',
        '',
        '规则2: 若有产生式 A → ...B',
        '       且 a ∈ LASTVT(B)',
        '       则 a ∈ LASTVT(A)',
    ])

    doc.add_paragraph()

    add_paragraph(doc, '优先关系表构建', bold_prefix='3. ')

    p = doc.add_paragraph()
    run = p.add_run('优先关系表的构建规则如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        '规则1: 若有产生式 A → ...ab... 或 A → ...aBb...',
        '       则 a = b',
        '',
        '规则2: 若有产生式 A → ...aB...',
        '       且 c ∈ FIRSTVT(B)',
        '       则 a <· c',
        '',
        '规则3: 若有产生式 A → ...Ba...',
        '       且 c ∈ LASTVT(B)',
        '       则 c ·> a',
    ])

    doc.add_paragraph()

    add_paragraph(doc, '移进-归约分析算法', bold_prefix='4. ')

    p = doc.add_paragraph()
    run = p.add_run('简单优先法的移进-归约分析算法流程如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        '算法：简单优先法语法分析',
        '输入：Token序列',
        '输出：语法分析结果（成功/失败）',
        '',
        '1. 初始化栈，将$压入栈中',
        '2. 循环执行以下步骤：',
        '   a. 找到栈顶第一个终结符a',
        '   b. 获取当前输入符号b',
        '   c. 比较a和b的优先关系：',
        '      - 若 a <· b 或 a = b，则移进，将b压入栈',
        '      - 若 a ·> b，则归约：',
        '        * 在栈中找到满足 ...X·>y 的最左符号X',
        '        * 将X归约为对应的非终结符',
        '      - 若 a = $ 且 b = $，则分析成功',
        '      - 否则，报告语法错误',
    ])

    doc.add_paragraph()

    add_note(doc, '【图表说明】请绘制一个简单优先法分析过程示例图，展示以下内容：')
    add_note(doc, '1. 使用一个表格展示分析过程，列包括：步骤、栈、当前输入、优先关系、动作')
    add_note(doc, '2. 以一个简单的表达式为例，如 id + id * id')
    add_note(doc, '3. 展示移进和归约的过程')

    doc.add_paragraph()

    # 3.2.3 语义分析器设计
    add_heading(doc, '3.2.3 语义分析器设计', 3)

    add_paragraph(doc, '设计原理', bold_prefix='1. ')

    p = doc.add_paragraph()
    run = p.add_run('语义分析器采用属性文法和语法制导翻译的方法进行设计。属性文法在文法的产生式中嵌入语义动作，当产生式被归约时，执行相应的语义动作。语义动作可以进行语义检查（如类型检查）、生成中间代码（如四元式）等操作。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_paragraph(doc, '功能实现', bold_prefix='2. ')

    p = doc.add_paragraph()
    run = p.add_run('语义分析器的主要功能包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    semantic_functions = [
        '（1）作用域管理：使用作用域栈管理变量的作用域，支持进入和退出作用域。',
        '（2）变量声明处理：在变量声明时，将变量信息添加到符号表中。',
        '（3）类型检查：检查表达式的类型是否正确，包括算术运算的类型兼容性、赋值语句的类型匹配等。',
        '（4）中间代码生成：生成四元式作为中间代码，每个四元式包含运算符、操作数1、操作数2、结果四个部分。',
    ]

    for func in semantic_functions:
        add_paragraph(doc, func, indent=True)

    add_paragraph(doc, '语义动作示例', bold_prefix='3. ')

    p = doc.add_paragraph()
    run = p.add_run('下面以赋值语句为例，说明语义动作的实现：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('产生式：S → id = E')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('语义动作：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'S → id = E',
        '{',
        '    // 1. 检查变量id是否已声明',
        '    if (!isDeclared(id)) {',
        '        reportError("变量未声明: " + id);',
        '    }',
        '    // 2. 检查类型兼容性',
        '    checkAssignment(id, E.type);',
        '    // 3. 生成四元式',
        '    emit("=", E.addr, "", id);',
        '}',
    ])

    doc.add_paragraph()

    # 3.2.4 代码生成器设计
    add_heading(doc, '3.2.4 代码生成器设计', 3)

    add_paragraph(doc, '设计原理', bold_prefix='1. ')

    p = doc.add_paragraph()
    run = p.add_run('代码生成器采用访问者模式（Visitor Pattern）遍历AST，为每种AST节点类型生成对应的VM指令。访问者模式将算法与对象结构分离，使得可以在不修改AST节点类的情况下定义新的操作。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_paragraph(doc, '功能实现', bold_prefix='2. ')

    p = doc.add_paragraph()
    run = p.add_run('代码生成器的主要功能包括：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    codegen_functions = [
        '（1）表达式代码生成：为算术运算、比较运算、逻辑运算生成VM指令。运算指令从栈中取操作数，将结果压入栈中。',
        '（2）控制流代码生成：为if-else、while等控制结构生成VM指令。使用标签和跳转指令实现控制流。',
        '（3）函数调用代码生成：为函数定义、函数调用、函数返回生成VM指令。处理参数传递、栈帧管理等。',
        '（4）变量访问代码生成：为变量读写生成push/pop指令，将变量值压入栈或从栈中弹出。',
    ]

    for func in codegen_functions:
        add_paragraph(doc, func, indent=True)

    add_paragraph(doc, '表达式代码生成算法', bold_prefix='3. ')

    p = doc.add_paragraph()
    run = p.add_run('表达式代码生成的算法流程如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        '算法：表达式代码生成',
        '输入：表达式AST节点',
        '输出：VM指令序列',
        '',
        'visitExpr(node):',
        '    switch (node.type):',
        '        case BinaryExpr:',
        '            // 1. 生成左操作数的代码',
        '            visitExpr(node.children[0])',
        '            // 2. 生成右操作数的代码',
        '            visitExpr(node.children[1])',
        '            // 3. 生成运算指令',
        '            emit(node.value)  // add, sub, mul, div...',
        '            break',
        '        case IdLeaf:',
        '            // 生成变量访问指令',
        '            emit("push", getSegment(node.value), getOffset(node.value))',
        '            break',
        '        case NumLeaf:',
        '            // 生成常量压栈指令',
        '            emit("push", "constant", node.value)',
        '            break',
    ])

    doc.add_paragraph()

    # ==================== 3.3 编码实现 ====================
    add_heading(doc, '3.3 编码实现', 2)

    # 3.3.1 开发环境搭建
    add_heading(doc, '3.3.1 开发环境搭建', 3)

    p = doc.add_paragraph()
    run = p.add_run('本项目使用CMake作为构建工具，CMakeLists.txt配置文件如下：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'cmake_minimum_required(VERSION 3.10)',
        'project(ShiXvn_Compiler)',
        '',
        'set(CMAKE_CXX_STANDARD 17)',
        '',
        '# 添加源文件',
        'set(SOURCES',
        '    main.cpp',
        '    src/Lexer.cpp',
        '    src/Grammar.cpp',
        '    src/PrecedenceParser.cpp',
        '    src/SemanticAnalyzer.cpp',
        '    src/AST.cpp',
        '    src/CodegenVisitor.cpp',
        '    src/SymbolTable.cpp',
        '    src/ErrorHandler.cpp',
        ')',
        '',
        '# 创建可执行文件',
        'add_executable(ShiXvn_Compiler ${SOURCES})',
        '',
        '# 添加头文件目录',
        'target_include_directories(ShiXvn_Compiler PRIVATE include)',
    ])

    doc.add_paragraph()

    # 3.3.2 核心代码实现
    add_heading(doc, '3.3.2 核心代码实现', 3)

    add_paragraph(doc, '词法分析器实现', bold_prefix='（1）')

    p = doc.add_paragraph()
    run = p.add_run('词法分析器的核心是nextToken()函数，它从源程序中读取下一个Token。以下是关键代码片段：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'Token Lexer::readIdentifierOrKeyword() {',
        '    int startLine = line_;',
        '    int startCol = column_;',
        '    string value;',
        '',
        '    // 读取字母、数字、下划线',
        '    while (isAlphaNumeric(currentChar())) {',
        '        value += currentChar();',
        '        advance();',
        '    }',
        '',
        '    // 检查是否为关键字',
        '    if (keywords_.count(value)) {',
        '        return {KEYWORD, value, startLine, startCol};',
        '    }',
        '',
        '    // 否则为标识符',
        '    return {IDENTIFIER, value, startLine, startCol};',
        '}',
    ])

    doc.add_paragraph()

    add_paragraph(doc, '语法分析器实现', bold_prefix='（2）')

    p = doc.add_paragraph()
    run = p.add_run('语法分析器的核心是parse()函数，它使用简单优先法对Token序列进行语法分析。以下是关键代码片段：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'bool PrecedenceParser::parse(const vector<Token>& tokens) {',
        '    // 初始化',
        '    push({"$", true, -1});  // 将$压入栈',
        '    size_t ip = 0;          // 输入指针',
        '',
        '    while (true) {',
        '        // 1. 找到栈顶终结符',
        '        int topIdx = findTopTerminal();',
        '        string a = stack_[topIdx].symbol;',
        '',
        '        // 2. 获取当前输入符号',
        '        string b = (ip < tokens.size()) ?',
        '                   tokenToTerminal(tokens[ip]) : "$";',
        '',
        '        // 3. 比较优先关系',
        '        string rel = grammar_.getRelation(a, b);',
        '',
        '        if (rel == "<" || rel == "=") {',
        '            // 移进',
        '            push({b, true, tokens[ip].line});',
        '            ip++;',
        '        } else if (rel == ">") {',
        '            // 归约',
        '            reduce(topIdx);',
        '        } else if (a == "$" && b == "$") {',
        '            // 接受',
        '            return true;',
        '        } else {',
        '            // 错误',
        '            reportError(tokens[ip]);',
        '            return false;',
        '        }',
        '    }',
        '}',
    ])

    doc.add_paragraph()

    add_paragraph(doc, '代码生成器实现', bold_prefix='（3）')

    p = doc.add_paragraph()
    run = p.add_run('代码生成器使用访问者模式遍历AST，为每种节点类型生成对应的VM指令。以下是关键代码片段：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'void CodegenVisitor::visitWhileStmt(const shared_ptr<ASTNode>& node) {',
        '    string L1 = sem_.newLabel();  // 循环开始标签',
        '    string L2 = sem_.newLabel();  // 循环结束标签',
        '',
        '    // 1. 输出循环开始标签',
        '    emit("label", L1, "", "");',
        '',
        '    // 2. 生成条件表达式代码',
        '    string condAddr = visitExpr(node->children[0]);',
        '',
        '    // 3. 条件为假时跳转到循环结束',
        '    emit("if-goto", L2, "", "");',
        '',
        '    // 4. 生成循环体代码',
        '    visit(node->children[1]);',
        '',
        '    // 5. 无条件跳转到循环开始',
        '    emit("goto", L1, "", "");',
        '',
        '    // 6. 输出循环结束标签',
        '    emit("label", L2, "", "");',
        '}',
    ])

    doc.add_paragraph()

    # ==================== 3.4 测试与验证 ====================
    add_heading(doc, '3.4 测试与验证', 2)

    # 3.4.1 测试策略
    add_heading(doc, '3.4.1 测试策略', 3)

    p = doc.add_paragraph()
    run = p.add_run('本项目的测试策略包括以下几个方面：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    test_strategies = [
        '（1）单元测试：对每个模块进行独立测试，验证模块功能的正确性。',
        '（2）集成测试：对多个模块进行组合测试，验证模块之间接口的正确性。',
        '（3）系统测试：对整个编译器进行端到端测试，验证编译器的整体功能。',
        '（4）回归测试：在修改代码后，重新运行之前的测试用例，确保修改没有引入新的错误。',
    ]

    for strategy in test_strategies:
        add_paragraph(doc, strategy, indent=True)

    # 3.4.2 测试用例设计
    add_heading(doc, '3.4.2 测试用例设计', 3)

    add_paragraph(doc, '测试用例1：简单算术运算', bold_prefix='')

    p = doc.add_paragraph()
    run = p.add_run('测试内容：计算1+2+...+10的和，验证变量声明、赋值、算术运算、循环等基本功能。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('源程序（test_flow.txt）：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        '// 计算 1+2+...+10 的和',
        'int sum ;',
        'int i ;',
        '',
        'sum = 0 ;',
        'i = 1 ;',
        '',
        'while ( i < 11 ) {',
        '    sum = sum + i ;',
        '    i = i + 1 ;',
        '} ;',
        '',
        '// sum 应该等于 55',
        'return sum ;',
    ])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('预期结果：RAM[256] = 55（1+2+...+10=55）')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('测试步骤：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    test_steps = [
        '步骤1：编译源程序',
        '  .\\ShiXvn_Compiler.exe test\\test_flow.txt',
        '  输出：output\\Codegen\\Main.vm',
        '',
        '步骤2：VMTranslator转换',
        '  .\\VMTranslator.exe output\\Codegen\\Main.vm output\\Codegen\\Main.asm',
        '  输出：output\\Codegen\\Main.asm',
        '',
        '步骤3：Assembler汇编',
        '  Assembler.bat output\\Codegen\\Main.asm',
        '  输出：output\\Codegen\\Main.hack',
        '',
        '步骤4：CPUEmulator执行',
        '  打开CPUEmulator，加载Main.hack，运行，查看RAM[256]',
    ]

    for step in test_steps:
        p = doc.add_paragraph()
        run = p.add_run(step)
        run.font.name = 'Courier New'
        run.font.size = Pt(10.5)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if step.startswith('  '):
            p.paragraph_format.left_indent = Pt(72)

    doc.add_paragraph()

    add_note(doc, '【图表说明】请插入测试用例1的运行结果截图，展示CPUEmulator的界面和RAM[256]=55的结果。')

    doc.add_paragraph()

    add_paragraph(doc, '测试用例2：冒泡排序', bold_prefix='')

    p = doc.add_paragraph()
    run = p.add_run('测试内容：实现冒泡排序算法，验证数组操作、嵌套循环、字符输出等高级功能。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('源程序（test_bubble.txt）：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        '// 冒泡排序测试',
        'int arr[10] ;',
        'int n ;',
        'int i ;',
        'int j ;',
        'int temp ;',
        '',
        '// 初始化数组（逆序排列）',
        'n = 10 ;',
        'set(arr, 0, 9) ;',
        'set(arr, 1, 8) ;',
        '// ... 更多初始化代码',
        '',
        '// 冒泡排序',
        'i = 0 ;',
        'while ( i < 9 ) {',
        '    j = 0 ;',
        '    while ( j < 9 ) {',
        '        if ( get(arr, j) > get(arr, j + 1) ) {',
        '            temp = get(arr, j) ;',
        '            set(arr, j, get(arr, j + 1)) ;',
        '            set(arr, j + 1, temp) ;',
        '        } ;',
        '        j = j + 1 ;',
        '    } ;',
        '    i = i + 1 ;',
        '} ;',
    ])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('预期输出：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_code(doc, [
        'Begin',
        '9 8 7 6 5 4 3 2 1 0',
        'End',
        '0 1 2 3 4 5 6 7 8 9',
    ])

    doc.add_paragraph()

    add_note(doc, '【图表说明】请插入测试用例2的运行结果截图，展示VMEmulator的界面和排序前后的输出结果。')

    doc.add_paragraph()

    # 3.4.3 测试结果与分析
    add_heading(doc, '3.4.3 测试结果与分析', 3)

    p = doc.add_paragraph()
    run = p.add_run('测试结果汇总如表3-5所示：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 测试结果表
    p = doc.add_paragraph()
    run = p.add_run('表3-5 测试结果汇总')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['测试用例', '测试内容', '预期结果', '实际结果', '测试状态']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    test_results = [
        ('测试1', '简单算术运算', 'RAM[256]=55', 'RAM[256]=55', '通过'),
        ('测试2', '冒泡排序', '排序正确', '排序正确', '通过'),
    ]

    for i, (col1, col2, col3, col4, col5) in enumerate(test_results, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        table.rows[i].cells[2].text = col3
        table.rows[i].cells[3].text = col4
        table.rows[i].cells[4].text = col5
        for j in range(5):
            for paragraph in table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('测试分析：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    analysis = [
        '（1）测试用例1验证了编译器的基本功能，包括变量声明、赋值、算术运算、循环等。测试结果表明，编译器能够正确地将源程序编译成VM指令，并在Hack虚拟机上得到正确的执行结果。',
        '（2）测试用例2验证了编译器的高级功能，包括数组操作、嵌套循环、字符输出等。测试结果表明，编译器能够正确处理复杂的程序结构，生成正确的VM指令。',
        '（3）通过两个测试用例的验证，证明编译器的各个模块（词法分析、语法分析、语义分析、代码生成）都能正确工作，编译器的整体功能满足设计要求。',
    ]

    for a in analysis:
        add_paragraph(doc, a, indent=True)

    # ==================== 3.5 本章小结 ====================
    add_heading(doc, '3.5 本章小结', 2)

    p = doc.add_paragraph()
    run = p.add_run('本章详细介绍了编译器的系统开发过程，包括总体设计、详细设计、编码实现和测试验证四个部分。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('在总体设计部分，介绍了编译器的整体架构设计、编译流程设计和数据结构设计。编译器采用模块化设计，各模块职责清晰，接口简洁。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('在详细设计部分，详细介绍了词法分析器、语法分析器、语义分析器和代码生成器的设计原理和实现方法。词法分析器采用有限自动机理论，语法分析器采用简单优先法，语义分析器采用属性文法和语法制导翻译，代码生成器采用访问者模式。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('在编码实现部分，介绍了开发环境搭建和核心代码实现。使用CMake作为构建工具，C++作为实现语言，展示了词法分析器、语法分析器和代码生成器的关键代码片段。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('在测试验证部分，介绍了测试策略、测试用例设计和测试结果。通过两个测试用例（简单算术运算和冒泡排序）的验证，证明编译器的各个模块都能正确工作，编译器的整体功能满足设计要求。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('通过本章的系统开发，成功实现了一个基于简单优先法的微型程序语言编译器，该编译器能够将mico C语言源程序编译成Hack虚拟机的VM指令，并在虚拟机上正确运行。')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存文档
    doc.save('D:\\c_cpp\\ShiXvn_Compiler\\report\\第三章_系统开发.docx')
    print('第三章生成成功！')

if __name__ == '__main__':
    create_chapter3()
