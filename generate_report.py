# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# 标题
title = doc.add_heading('周报', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 日期信息
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run('项目名称：基于简单优先文法的 Micro-C 编译器')
date_run.font.size = Pt(12)

date_para2 = doc.add_paragraph()
date_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run2 = date_para2.add_run('报告周期：2026年6月20日 - 2026年6月27日')
date_run2.font.size = Pt(11)

doc.add_paragraph()

# ==================== 本周工作总结 ====================
doc.add_heading('一、本周工作总结', level=1)

# 1.1 编译器架构重构
doc.add_heading('1. 编译器架构重构', level=2)

doc.add_heading('（1）AST 基础设施搭建', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('新建 ')
p.add_run('include/AST.h').bold = True
p.add_run(' 和 ')
p.add_run('src/AST.cpp').bold = True
p.add_run('，定义完整的抽象语法树节点类型')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('支持节点类型：Program、VarDecl、AssignStmt、IfStmt、WhileStmt、BinaryExpr、CallExpr、LValue 等')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('实现 dumpAST() 递归打印工具，输出到 output/Parser/ast.txt')

doc.add_heading('（2）解耦 AST 构建与代码生成', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('修改 PrecedenceParser，在归约时同步构建 AST 节点')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('新增 buildASTForProduction() 方法，与原有 semanticAction() 并行执行')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('AST 根节点可通过 parser.getAST() 获取')

doc.add_heading('（3）符号表升级为作用域栈', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('将 SemanticAnalyzer 中的扁平 map 升级为 vector<Scope> 作用域栈')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('实现 enterScope() / exitScope() 支持块级作用域')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('实现 lookup() 从内到外查找变量')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('新增函数表 FuncEntry，支持函数签名校验')

doc.add_heading('（4）CodegenVisitor 准备', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('新建 include/CodegenVisitor.h 和 src/CodegenVisitor.cpp')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('实现访问者模式，用于后续从 AST 生成四元式')

# 1.2 文法扩展
doc.add_heading('2. 文法扩展（SPG 限制内）', level=2)

doc.add_heading('（1）数组声明支持', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('新增产生式：S → int id [ num ]、S → char id [ num ]、S → double id [ num ]')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('数组声明时自动调用 Memory.alloc 分配堆内存')

doc.add_heading('（2）if 语句支持（无 else）', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('新增产生式：S → if ( E ) { P }')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('支持单分支 if 语句，无需 else 分支')

doc.add_heading('（3）优先关系表扩展', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('添加 } < else 处理 if-else 结构')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('添加 } > ; 处理块后分号')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('添加 $ < ; 处理多语句程序')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('添加 { < ; 处理块内多语句')

# 1.3 数据结构与内置函数
doc.add_heading('3. 数据结构与内置函数', level=2)

doc.add_heading('（1）结构体支持', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('结构体变量声明时自动调用 Memory.alloc 分配堆内存')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('实现 set(struct, member, value) 和 get(struct, member) 函数')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('成员名自动转换为偏移量（基于结构体定义）')

doc.add_heading('（2）数组支持', level=3)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('数组声明时自动分配堆内存')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('set(arr, index, value) 和 get(arr, index) 支持变量作为索引')

doc.add_heading('（3）内置函数完善', level=3)

# 内置函数表格
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['函数', 'Jack VM 映射', '说明']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['read()', 'String.new + Keyboard.readInt 1', '从键盘读取整数'],
    ['write(expr)', 'Output.printInt', '输出整数'],
    ['writeChar(c)', 'Output.printChar', '输出字符'],
    ['writeString(s)', 'Output.printString', '输出字符串'],
    ['println()', 'Output.println', '输出换行'],
    ['set(base, member, value)', '内联 5 条指令', '写入内存'],
    ['get(base, member)', '内联 3 条指令', '读取内存'],
]

for i, row_data in enumerate(data):
    for j, cell_text in enumerate(row_data):
        table.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph()

# 1.4 词法分析器修复
doc.add_heading('4. 词法分析器修复', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('恢复 readNumber() 中的浮点数处理代码')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('3.14 现在正确识别为单个 NUMBER token')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('浮点数在代码生成时截断为整数（Jack VM 限制）')

# 1.5 测试用例
doc.add_heading('5. 测试用例', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('test_bubble.txt').bold = True
p.add_run(' — 冒泡排序（双重循环 + 数组 + 条件判断）')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('test_char.txt').bold = True
p.add_run(' — 字符输出测试')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('test_pop.txt').bold = True
p.add_run(' — 单次冒泡遍历测试')

doc.add_paragraph()

# ==================== 下周工作计划 ====================
doc.add_heading('二、下周工作计划', level=1)

doc.add_heading('1. 完善 CodegenVisitor', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('将 CodegenVisitor 接入主流程，替代原有的内联代码生成')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('实现从 AST 到四元式的完整遍历')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('移除 PrecedenceParser 中的 semanticAction() 依赖')

doc.add_heading('2. 文法扩展', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('支持 for 循环语法')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('支持 do-while 循环')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('支持多维数组 int arr[3][4]')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('支持结构体嵌套 struct Line { struct Point p1 ; struct Point p2 ; }')

doc.add_heading('3. 类型系统增强', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('实现完整的类型检查（函数参数类型匹配）')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('支持隐式类型转换规则')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('添加数组越界检查（编译时）')

doc.add_heading('4. 代码优化', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('优化临时变量使用（减少 temp 溢出到 local 段）')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('实现常量折叠（编译时计算 3 + 4 → 7）')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('优化控制流代码（减少冗余跳转）')

doc.add_heading('5. 文档完善', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('更新 README.md 中的文法定义')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('编写编译器使用指南')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('整理测试用例集')

# 保存文件
output_path = r'D:\c_cpp\ShiXvn_Compiler\report\周报_2026-06-27.docx'
doc.save(output_path)
print(f'周报已保存到: {output_path}')
